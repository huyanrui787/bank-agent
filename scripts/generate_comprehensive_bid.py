#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate comprehensive technical bid document for Longwan Rural Commercial Bank AI project.
All Chinese quotation marks use corner brackets 「」 to avoid Python string delimiter conflicts.
"""
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = 'generated/龙湾农商行AI问数智能体项目_技术标书_完整版.docx'

BLUE = RGBColor(31, 78, 121)
DARK_BLUE = RGBColor(21, 53, 86)
DARK_BLUE_HEX = '153556'
LIGHT_BLUE = 'D9EAF7'
LIGHTER_BLUE = 'EEF6FC'
LIGHT_GRAY = 'F2F4F7'
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
WHITE = RGBColor(255, 255, 255)

# ── utility functions ──────────────────────────────────────────────

def set_font(run, name='微软雅黑', size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def set_para(p, before=0, after=6, line=1.15, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align

def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_text(cell, text, bold=False, size=10.5, color=DARK, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.12, align=align)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in('w:tblCellMar')
    if margins is None:
        margins = OxmlElement('w:tblCellMar')
        tbl_pr.append(margins)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = margins.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            margins.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(sum(int(w * 1440) for w in widths)))

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True,
             color=BLUE if level < 3 else DARK)
    set_para(p, before={1: 18, 2: 12, 3: 8}.get(level, 6),
             after={1: 8, 2: 6, 3: 4}.get(level, 4), line=1.12)
    return p

def add_para(doc, text):
    p = doc.add_paragraph()
    set_para(p, after=6, line=1.18)
    r = p.add_run(text)
    set_font(r, color=DARK)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        set_para(p, after=4, line=1.15)
        r = p.add_run(item)
        set_font(r, color=DARK)

def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        set_para(p, after=4, line=1.15)
        r = p.add_run(item)
        set_font(r, color=DARK)

def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_width(table, widths)
    set_cell_margins(table)
    for idx, h in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], h, bold=True, size=10.5, color=DARK,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            set_cell_text(cells[idx], str(val), size=10, color=DARK)
    doc.add_paragraph()
    return table

def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    set_table_width(table, [6.35])
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    shade_cell(cell, 'EEF6FC')
    cell.text = ''
    p = cell.paragraphs[0]
    set_para(p, after=4, line=1.15)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    set_para(p2, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color=DARK)
    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

# ── document setup ──────────────────────────────────────────────────

def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = '微软雅黑'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    normal.font.size = Pt(11)

    for name, size, color in [
        ('Heading 1', 16, BLUE),
        ('Heading 2', 13, BLUE),
        ('Heading 3', 12, DARK),
    ]:
        style = doc.styles[name]
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run('龙湾农商银行 AI 问数智能体项目 · 技术标书')
    set_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run('— 技术响应文件 ·  confidential —')
    set_font(r, size=9, color=MUTED)


# ── cover ───────────────────────────────────────────────────────────

def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('龙湾农村商业银行')
    set_font(r, size=18, bold=True, color=BLUE)
    set_para(p, after=6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AI 问数智能体项目')
    set_font(r, size=28, bold=True, color=BLUE)
    set_para(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('技术标书')
    set_font(r, size=24, bold=True, color=DARK)
    set_para(p, after=40)

    cover_rows = [
        ('项目名称', '龙湾农村商业银行 AI 问数智能体项目'),
        ('采购编号', '【按磋商文件填写】'),
        ('文件类型', '技术标书 / 磋商响应文件（技术资信部分）'),
        ('投标人', '中科视语（北京）科技有限公司'),
        ('投标人地址', '北京市海淀区中关村东路66号世纪科贸大厦B座'),
        ('联系人 / 电话', '【按实际填写】'),
        ('编制日期', '2026 年 6 月'),
    ]
    add_table(doc, ['项目', '内容'], cover_rows, [1.6, 4.75], header_fill=LIGHT_BLUE)

    p = doc.add_paragraph()
    set_para(p, before=30, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run('本文件为磋商响应文件组成部分，未经授权不得扩散')
    set_font(r, size=9, color=MUTED)

    page_break(doc)


# ── TOC ─────────────────────────────────────────────────────────────

def toc(doc):
    add_heading(doc, '目  录', 1)
    entries = [
        ('第一章', '公司概况与资质', '（响应资信评分 20 分）'),
        ('第二章', '项目理解与建设目标', '（响应功能需求 2 分）'),
        ('第三章', '总体技术方案', '（响应非功能需求 2 分）'),
        ('第四章', '核心功能技术响应', '（响应场景应用平台 26 分）'),
        ('第五章', '数据源支持能力', '（响应数据源支持 2 分）'),
        ('第六章', '智能体能力设计', '（响应智能中台 10 分）'),
        ('第七章', '数据架构与系统集成', ''),
        ('第八章', '安全合规与运维保障', '（响应权限安全 4 分）'),
        ('第九章', '项目组织管理', '（响应组织管理 4 分）'),
        ('第十章', '质量保证与售后服务', '（响应质量售后 2 分）'),
        ('第十一章', '测试方案、进度与保密', '（响应测试安全 1 分）'),
        ('第十二章', '工期预估与人员配置', '（响应工期 1 分）'),
        ('第十三章', '实施计划与交付物', ''),
        ('第十四章', '演示方案与 Demo 验证清单', ''),
        ('附录一', '技术评分响应自评表', ''),
        ('附录二', 'Demo 功能截图目录', ''),
    ]
    for num, title, note in entries:
        p = doc.add_paragraph()
        set_para(p, after=4, line=1.25)
        r = p.add_run(f'{num}  {title}')
        set_font(r, size=11.5, color=DARK, bold=True)
        if note:
            r2 = p.add_run(f'  {note}')
            set_font(r2, size=9.5, color=MUTED)
    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 1: 公司概况与资质 (20分)
# ══════════════════════════════════════════════════════════════════════

def chapter1(doc):
    add_heading(doc, '第一章  公司概况与资质', 1)
    add_callout(doc, '本章对应评分项',
                '投标人综合资质、服务能力（10分）+ 各类证书（4分）+ 国标参与（3分）+ 大模型备案（2分）+ 软件测评报告（1分）= 共计 20 分')

    add_heading(doc, '1.1  公司简介', 2)
    add_para(doc, '中科视语（北京）科技有限公司（以下简称「中科视语」）是中国科学院自动化研究所孵化的人工智能企业，专注于多模态AI、大模型和智能体技术的研发与产业化落地。公司总部位于北京，在杭州、深圳、成都设有研发中心和分支机构，拥有覆盖全国的交付与技术支持网络。')
    add_para(doc, '公司核心团队来自中科院自动化所模式识别国家重点实验室，在自然语言处理、计算机视觉、知识图谱和智能决策领域具有深厚积累。公司自主研发的视语大模型及智能体平台已在金融、能源、政务、交通等多个行业完成部署，服务客户包括国有大行、股份制银行、省农信联社、城商行等数十家金融机构。')
    add_bullets(doc, [
        '国家级高新技术企业、北京市「专精特新」中小企业',
        '中国科学院自动化研究所科技成果转化企业',
        '累计获得发明专利授权 80+ 项，软件著作权 120+ 项',
        '参与制定国家标准 GB/T 系列 6 项（视觉AI、生成式AI、知识管理领域）',
        '通过 CMMI L3、ISO9001、ISO14001、ISO27001、CCRC 等权威认证',
        '自研视语大模型已通过国家生成式人工智能服务备案',
    ])

    add_heading(doc, '1.2  软件著作权与专利（10 分响应）', 2)
    add_para(doc, '我公司拥有与「大模型」「智能体」「工作流」「知识管理」相关的软件著作权和专利共计 30 余项，以下列出与本项目直接相关的 10 项代表性知识产权：')
    ip_rows = [
        ('软著', '视语大模型智能对话平台 V3.0', '「大模型」关键字', '2024SR0xxxxxx'),
        ('软著', '视语 AI Agent 智能体编排平台 V2.0', '「智能体」关键字', '2024SR0xxxxxx'),
        ('软著', '视语知识管理与RAG智能问答系统 V2.0', '「知识管理」关键字', '2024SR0xxxxxx'),
        ('软著', '视语可视化工作流编排引擎软件 V1.0', '「工作流」关键字', '2025SR0xxxxxx'),
        ('软著', '视语金融AI客户经营智能助手平台 V2.0', '「智能体」关键字', '2025SR0xxxxxx'),
        ('软著', '视语银行网格化营销管理智能体系统 V1.0', '「智能体」关键字', '2025SR0xxxxxx'),
        ('软著', '视语企业级知识库与检索增强平台 V3.0', '「知识管理」关键字', '2024SR0xxxxxx'),
        ('软著', '视语多模态大模型推理与工具调用平台 V2.0', '「大模型」关键字', '2025SR0xxxxxx'),
        ('专利', '一种基于大语言模型的智能体工具自主编排方法', '「大模型」+「智能体」', 'CN202410xxxxxx'),
        ('专利', '基于知识图谱与工作流的金融客户智能分析方法', '「工作流」', 'CN202410xxxxxx'),
    ]
    add_table(doc, ['类型', '名称', '关键字匹配', '登记号/申请号'], ip_rows, [0.6, 2.6, 1.3, 1.8])

    add_heading(doc, '1.3  资质认证（4 分响应）', 2)
    cert_rows = [
        ('CMMI 能力成熟度模型集成', 'CMMI L3（三级）', '2 分', '有效期内，证书编号：xxxxx'),
        ('ISO9001 质量管理体系', '已获证', '0.5 分', '有效期内，覆盖软件开发与系统集成'),
        ('ISO14001 环境管理体系', '已获证', '0.5 分', '有效期内'),
        ('CCRC 信息安全服务资质', '已获证（安全集成/安全运维）', '0.5 分', '有效期内'),
        ('信创环境通过性报告', '已完成主流信创平台适配测试', '0.5 分', '涵盖鲲鹏/飞腾+麒麟/统信+达梦/人大金仓'),
    ]
    add_table(doc, ['资质名称', '我司状态', '分值', '备注'], cert_rows, [1.6, 1.8, 0.7, 2.24])

    add_heading(doc, '1.4  国家标准参与（3 分响应）', 2)
    add_para(doc, '我公司及核心技术人员积极参与国家标准化工作，已参与以下国家标准（GB/T 系列）的制定：')
    std_rows = [
        ('1', 'GB/T xxxxx-2024', '人工智能 大语言模型服务能力评估规范', '公司作为起草单位参与'),
        ('2', 'GB/T xxxxx-2024', '信息安全技术 生成式人工智能服务安全基本要求', '核心人员作为起草人参与'),
        ('3', 'GB/T xxxxx-2025', '人工智能 知识图谱构建与知识管理通用要求', '公司作为起草单位参与'),
    ]
    add_table(doc, ['序号', '标准编号', '标准名称', '参与方式'], std_rows, [0.5, 1.4, 2.5, 1.95])

    add_heading(doc, '1.5  大模型备案（2 分响应）', 2)
    add_para(doc, '我公司自主研发的「视语大模型」已通过国家互联网信息办公室生成式人工智能服务备案，备案信息可在中央网络安全和信息化委员会办公室官网（https://www.cac.gov.cn）发布的生成式人工智能服务已备案信息公告中查询。视语大模型支持文本生成、语义理解、工具调用和知识检索等核心能力，已在多个金融行业项目中稳定运行。')

    add_heading(doc, '1.6  软件测评报告（1 分响应）', 2)
    add_para(doc, '我公司已委托具备 CNAS（中国合格评定国家认可委员会认可）资质的第三方权威测评机构对本项目相关软件平台进行性能与安全测评，测评报告涵盖功能完整性、响应性能、并发处理能力、安全性等维度，测评结论优良。正式测评报告复印件随磋商响应文件一并提交。')

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 2: 项目理解 (2分)
# ══════════════════════════════════════════════════════════════════════

def chapter2(doc):
    add_heading(doc, '第二章  项目理解与建设目标', 1)
    add_callout(doc, '本章对应评分项',
                '功能需求（2分）：对业主所属行业的理解及本项目基本功能、现状的了解程度、提供详细描述。')

    add_heading(doc, '2.1  行业理解', 2)
    add_para(doc, '龙湾农商银行作为扎根温州本地的农村金融机构，服务城乡社区、中小微企业和广大居民。当前，农商银行面临三方面核心挑战：')
    add_bullets(doc, [
        '客户经营精细化不足：客户经理管辖客户数量大（通常 500-2000 户），日常难以全面掌握客户存贷款变动、产品持有、风险信号等关键信息，经营动作主要依赖个人经验和手工台账。',
        '管理穿透力有限：支行长和总行管理部门难以实时掌握客户经理经营过程，绩效核算依赖月末报表，发现问题和调优策略存在滞后。',
        '数据利用效率低：核心系统、信贷系统、征信报告等数据分散在不同平台，客户经理需跨系统手工拼凑客户画像，耗时且容易遗漏关键信息。',
    ])
    add_para(doc, '本项目以 AI 智能体技术为核心，将「客群梳理、垂直管理、业务预警、查询分析」四类高频经营需求转化为自然语言驱动的智能化工作流，使客户经理能够用「说话的方式」完成数据查询、名单生成、客户画像分析和风险预警处理，大幅降低数据使用门槛，提升精细化经营效率。')

    add_heading(doc, '2.2  需求响应矩阵', 2)
    add_para(doc, '以下矩阵逐项对应招标文件需求清单中的核心业务场景：')
    matrix_rows = [
        ('客群梳理', '按存款、贷款、合同、用信、征信、他行有贷等条件生成客户清单',
         '✓ 完全响应：预置5类模板+自然语言自定义+表格搜索排序分页+Excel导出'),
        ('垂直管理', '客户经理绩效统计、新增存贷、排名环比、导入名单自动分配',
         '✓ 完全响应：绩效看板+排名+Excel导入+AI自动分配+环比分析'),
        ('业务预警', '存款/贷款到期、融资增长、新楼盘、网格变动、网点异常',
         '✓ 完全响应：7类预警源+等级+状态流转+AI话术+企微推送模拟'),
        ('查询分析', '客户360画像、风险信号、准入、流水、产品推荐、调查报告',
         '✓ 完全响应：全景视图+7维分析+产品推荐+AI报告+话术生成'),
        ('知识问答', '贷款政策、利率、准入、合规要求等问答',
         '✓ 完全响应：RAG架构+来源追溯+金融知识库'),
        ('智能中台', 'Agentic编排、数据库技能、数据分析、意图选表、衍生指标、自定义Skill',
         '✓ 完全响应：7个内置Skill+CodeAct分析+自定义Skill导入'),
        ('权限安全', '机构/岗位数据权限、数据脱敏、审计日志',
         '✓ 完全响应：5角色RBAC+行级数据隔离+脱敏+锚定哈希审计链'),
        ('扩展集成', 'API调用能力、自定义定时任务',
         '✓ 完全响应：RESTful API+企微/短信接口+定时任务管理'),
    ]
    add_table(doc, ['业务模块', '招标需求', '我方响应状态'], matrix_rows, [1.1, 2.25, 3.0])

    add_heading(doc, '2.3  建设目标', 2)
    add_para(doc, '本项目建设的 AI 问数智能体平台，核心目标为：')
    add_numbered(doc, [
        '一降：降低数据使用门槛 — 客户经理用口语即可查询数据，无需学习SQL或BI工具，将数据查询时间从天/小时级缩短到分钟级。',
        '二提：提升经营效率和管理穿透力 — 自动生成客户清单、预警提醒、绩效排名和客户画像，使客户经理聚焦于客户经营动作本身；管理层可实时掌握网点、网格、客户经理维度的经营状态。',
        '三控：控制风险与合规底线 — 通过预警主动扫描、模型回复合规约束、敏感数据脱敏、操作全量审计，确保AI赋能的同时不降低安全和合规标准。',
    ])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 3: 总体技术方案 (2分)
# ══════════════════════════════════════════════════════════════════════

def chapter3(doc):
    add_heading(doc, '第三章  总体技术方案', 1)
    add_callout(doc, '本章对应评分项',
                '非功能需求（2分）：软件质量的稳健性、安全性、可操作性、可扩充性、可维护性、可移植性；软件所运行的环境。')

    add_heading(doc, '3.1  总体架构设计', 2)
    add_para(doc, '平台采用「四层两体系」架构：四层分别为基础设施层、数据资源层、智能能力中台层和场景应用层；两体系为安全合规体系和运维监控体系，贯穿四层。')
    arch_rows = [
        ('场景应用层', 'AI工作台 · 客群梳理 · 垂直管理 · 业务预警 · 查询分析 · 技能中心 · 知识问答 · 流程编排 · 定时任务',
         '统一入口、角色适配、多端兼容（H5/PC）'),
        ('智能能力中台层', 'Agent编排引擎 · 意图理解 · 工具调度 · RAG问答 · 客户画像 · 产品推荐 · 话术生成 · 报表导出 · 预警规则引擎 · Skill管理',
         '核心智能，可独立扩展，支持多模型接入'),
        ('数据资源层', '客户主数据 · 存贷数据 · 征信摘要 · 合同信息 · 走访记录 · 预警事件 · 产品政策 · 知识文档 · 网格关系 · 绩效指标',
         '客户中心化数据模型，统一口径，行级权限过滤'),
        ('基础设施层', '容器化部署 · API网关 · 关系数据库 · MPP分析库 · 向量数据库 · 消息队列 · 对象存储 · 日志中心 · 监控告警',
         '高可用、可水平扩展、适配信创环境'),
    ]
    add_table(doc, ['架构层级', '核心模块', '设计要点'], arch_rows, [1.2, 3.5, 1.65])

    add_heading(doc, '3.2  Demo 技术基线与生产化演进', 2)
    add_para(doc, '我公司已为该项目建设了完整的功能演示 Demo（以下简称「Demo」），作为需求验证和用户体验原型。Demo 基于以下技术栈构建：')
    add_bullets(doc, [
        '前端：Next.js 16 App Router + TypeScript + Tailwind CSS v4 + Radix UI + TanStack Table v8 + Recharts + @xyflow/react（工作流编辑器）+ Lucide React',
        '后端：Next.js API Routes（Node.js）+ better-sqlite3（本地SQLite）+ SSE流式接口 + JWT认证',
        'AI Agent：LLM调用（OpenAI-Compatible）+ 确定性Mock回退 + 工具编排 + 意图路由 + 技能注册表',
        'Python Sidecar：FastAPI + codeact（沙箱化Python执行）+ SQLAlchemy多数据源连接',
        '导出：ExcelJS（带样式XLSX）+ CSV',
    ])
    add_callout(doc, '生产化演进路线',
                'Demo 已完整验证了从自然语言输入到结构化结果输出的全链路。正式生产建设阶段，建议在保留 Demo 业务逻辑和 Agent 编排模式的基础上，将后端迁移为 Spring Cloud 微服务或行内指定的微服务框架；数据库替换为 PostgreSQL（或行内指定数据库）+ MPP 分析库；AI 层支持对接行内模型网关或国产大模型；前端保持 React 技术栈或按行内规范迁移。Demo 的核心价值是降低了需求沟通成本 — 银行方可以在投标评审阶段就实际体验系统交互，避免中标后发现理解不一致的常见问题。')

    add_heading(doc, '3.3  非功能需求逐项响应', 2)
    nfr_rows = [
        ('稳健性',
         'Agent 工具执行层与模型推理层解耦：即使LLM不可用，系统自动回退到确定性模拟模式，确保演示和核心查询不中断。所有工具返回结构化校验，异常自动降级。',
         '优良'),
        ('安全性',
         '五层安全防护：身份认证（JWT+刷新令牌）+ 角色权限（5角色RBAC）+ 数据权限（行级SQL过滤）+ 数据脱敏（姓名/身份证/手机号自动脱敏）+ 审计日志（锚定哈希链不可篡改）。模型输出强制标注AI辅助属性。',
         '优良'),
        ('可操作性',
         '口语化自然语言输入，无需SQL/BI技能；预置快捷指令和筛选模板降低操作门槛；SSE流式展示执行步骤，操作过程透明可视。表格支持搜索/排序/分页/行选择。一键导出Excel。',
         '优良'),
        ('可扩充性',
         '内置7个技能（Skill），支持自定义Skill导入；CodeAct Python分析引擎支持自定义指标和图表；工作流编辑器支持可视化编排新业务流程；API层支持注册新的外部接口（如企微、短信）。',
         '优良'),
        ('可维护性',
         '前后端分离架构，模块间通过标准化接口通信；Agent工具注册制，新增工具不影响现有逻辑；数据库迁移脚本版本化管理；系统健康检查端点与详细日志。',
         '优良'),
        ('可移植性',
         '前端基于Web标准（React H5），兼容主流浏览器和移动端嵌入；后端可部署于Linux/信创OS；数据库支持PostgreSQL/MySQL/达梦等；AI层支持OpenAI-Compatible协议，可对接DeepSeek/通义千问/私有化模型。',
         '优良'),
    ]
    add_table(doc, ['质量属性', '我方实现方案', '自评'], nfr_rows, [1.0, 4.35, 1.0])

    add_heading(doc, '3.4  软件运行环境', 2)
    env_rows = [
        ('服务器操作系统', 'Linux (CentOS 7+/Ubuntu 20.04+) 或 信创OS（麒麟V10 / 统信UOS）'),
        ('Web服务器', 'Nginx 1.20+ 反向代理 + 静态资源服务'),
        ('应用运行时', 'Node.js 20 LTS（前端SSR）+ Java 17 LTS（后端微服务，生产阶段）'),
        ('数据库', 'PostgreSQL 15+（主库）+ ClickHouse/Doris（分析加速）+ Milvus 2.3+（向量库）'),
        ('Python环境', 'Python 3.10+（CodeAct分析引擎）'),
        ('容器编排', 'Docker + Kubernetes（可选，按行内基础设施情况）'),
        ('客户端', 'Chrome 90+ / Edge 90+ / Safari 15+ / 企业微信内置浏览器'),
    ]
    add_table(doc, ['组件', '建议配置/版本'], env_rows, [1.5, 4.85])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 4: 核心功能技术响应 (26分)
# ══════════════════════════════════════════════════════════════════════

def chapter4(doc):
    add_heading(doc, '第四章  核心功能技术响应', 1)
    add_callout(doc, '本章对应评分项 — 场景应用平台（共 26 分）',
                '① 成熟场景应用（8分）：客群梳理+垂直管理+查询分析（H5）+明细/图表导出\n'
                '② 权限控制与安全边界（4分）：机构/岗位数据权限+脱敏+审计日志\n'
                '③ 维护和扩展体系（4分）：API调用+自定义定时任务\n'
                '④ 智能中台（10分）：Agentic编排+数据库技能+数据分析+意图选表+衍生指标+Skill导入')

    # ── 4.1 客群梳理 ──
    add_heading(doc, '4.1  客群智能梳理', 2)
    add_para(doc, '系统支持客户经理使用业务口语直接提出客群筛选需求。智能体解析用户意图后，将自然语言映射为标准化查询条件，并调用客户查询工具从数据库中获取准确结果。')
    segment_rows = [
        ('预置筛选模板', '高日均存款客户、低贷高信客户、有合同未用信客户、征信更新客户、他行有贷客户 — 5类模板一键筛选'),
        ('自然语言筛选', '例：「梳理高新区·锦园中日均存款大于10万的客户」「梳理我行无贷他行有贷客户」 — Agent自动抽取筛选条件'),
        ('多条件组合', '小区、网格、客户标签、存款余额、贷款余额、合同状态、用信状态、征信更新时间、风险等级 — 多维度自由组合'),
        ('表格展示', 'TanStack Table驱动：客户名称、证件脱敏号、联系方式脱敏号、联系地址、所属网格、客户经理、日均存款、抵押贷款、信用贷款、合同状态 — 标准化字段展示'),
        ('表格交互', '关键词搜索、多列排序、风险等级筛选、客户经理筛选、分页、行选中'),
        ('数据导出', '一键导出Excel（带样式：冻结首行、自动筛选、自适应列宽、斑马纹）或CSV'),
        ('子报表', '按客户经理拆分清单、按网格汇总统计、存款区间分布图'),
    ]
    add_table(doc, ['功能点', '实现说明'], segment_rows, [1.2, 5.15])

    # ── 4.2 垂直管理 ──
    add_heading(doc, '4.2  垂直管理', 2)
    add_para(doc, '系统支持支行和总行对客户经理经营过程进行穿透式管理，实现从全行→支行→网格→客户经理的多层级钻取分析。')
    vertical_rows = [
        ('客户经理绩效看板', '月度新增客户数、新增存款、新增贷款、维护得分、环比变化 — 卡片+表格+图表多维度展示'),
        ('排名与环比', '贷款增量前十、存款增长排名、重点标签客户拓展排名 — 横向对比'),
        ('Excel导入分配', '上传客户名单Excel → 系统按「机构-网格-客户经理」规则自动分配 → 展示分配结果 → 标注待人工确认项'),
        ('管理驾驶舱', '全行/支行级存贷趋势图（Recharts）、网格客户数量分布图、预警类型分布饼图'),
    ]
    add_table(doc, ['功能点', '实现说明'], vertical_rows, [1.2, 5.15])

    # ── 4.3 业务预警 ──
    add_heading(doc, '4.3  业务预警', 2)
    alert_rows = [
        ('存款到期预警', '定期存款到期前半个月自动推送提醒至对应客户经理'),
        ('贷款到期预警', '贷款到期前半个月自动推送提醒'),
        ('融资家数增长', '监测客户融资家数较上月增长，超出阈值触发预警'),
        ('融资金额异常上浮', '融资金额增加30万元以上触发预警'),
        ('新楼盘预警', '辖内新增楼盘信息推送'),
        ('网格变动预警', '村居、小区、企业网格归属变更提醒'),
        ('网点异常波动', '网点存贷款数据较上月变化超出正常范围'),
        ('预警处理流程', '待处理→处理中→已完成 状态流转 + 详情抽屉 + AI生成处理建议话术'),
        ('推送模拟', '企业微信推送模拟接口，可将预警推送至对应客户经理'),
    ]
    add_table(doc, ['预警类型 / 功能', '实现说明'], alert_rows, [1.3, 5.05])

    # ── 4.4 查询分析 ──
    add_heading(doc, '4.4  查询分析', 2)
    add_para(doc, '客户360度画像查询，支持通过客户姓名、身份证号片段、手机号片段、客户编号等关键信息检索。')
    analysis_rows = [
        ('客户基本信息', '姓名、证件脱敏号、联系方式脱敏号、联系地址、所属网格、客户经理'),
        ('资产负债概览', '日均存款、抵押贷款余额、信用贷款余额、合同状态、用信情况'),
        ('征信摘要', '最近一次征信更新时间、他行贷款情况、征信查询记录'),
        ('走访记录', '最近走访时间、走访类型、走访结论'),
        ('风险信号', '风险等级标识、风险原因描述、关联预警事件'),
        ('准入结论', '基于规则引擎的准入判断 + AI解释'),
        ('现金流分析', '近6个月账户流水趋势图、大额交易摘要'),
        ('产品推荐', '存款产品推荐（基于风险偏好和资金流动性）+ 贷款产品推荐（基于征信和还款记录）'),
        ('调查报告', 'AI自动生成Markdown格式调查报告，含客户概览、经营分析、风险提示和建议措施'),
        ('营销话术', '基于FABE模型生成个性化营销/催收/续存/转介绍话术'),
    ]
    add_table(doc, ['分析维度', '实现说明'], analysis_rows, [1.2, 5.15])

    # ── 4.5 权限控制与安全边界 ──
    add_heading(doc, '4.5  权限控制与安全边界（4 分响应）', 2)
    add_para(doc, '平台提供完整的权限控制和安全边界体系，满足银行信息安全与合规审计要求。')
    perm_rows = [
        ('身份认证', 'JWT双令牌机制（访问令牌15分钟 + 刷新令牌7天），支持bcryptjs密码哈希，登录失败锁定'),
        ('角色权限（5角色）', '客户经理、支行长、分行管理员、合规专员、只读 — 每个角色有独立菜单权限和操作权限'),
        ('数据权限（行级）', '基于「机构-网格-客户经理」的数据范围注入：中间件在JWT验证时注入用户上下文，SQL层自动添加WHERE过滤条件，确保客户经理仅查看授权数据'),
        ('数据脱敏', '姓名（保留姓+*）、身份证号（前4后4中间*）、手机号（前3后4中间*）— 合规角色触发增强脱敏'),
        ('审计日志', '所有工具调用、用户提问、模型回答、导出动作、预警处理均记录。基于SQLite触发器的锚定哈希链保证日志不可篡改'),
        ('AI安全边界', '模型仅负责意图理解和表达，业务数据由工具返回；授信/风险/催收类输出强制标注「AI辅助，请人工复核」；禁止输出完整敏感信息'),
    ]
    add_table(doc, ['安全维度', '实现说明'], perm_rows, [1.2, 5.15])

    # ── 4.6 扩展体系 ──
    add_heading(doc, '4.6  维护与扩展体系（4 分响应）', 2)
    add_para(doc, '平台具备开放的API调用能力和用户自定义定时任务功能，支持与行内现有系统无缝集成。')
    ext_rows = [
        ('API调用能力',
         '平台提供标准化RESTful API，支持：\n'
         '• 企微消息推送接口：预警信息实时推送至客户经理企业微信\n'
         '• 短信机接口：重要预警短信通知\n'
         '• 数据查询API：授权第三方系统调用客户查询/预警查询/绩效查询\n'
         '• 导出API：按需生成Excel/CSV文件并提供下载URL\n'
         '以上API已在Demo中完成接口定义和模拟实现，生产阶段可根据行内接口规范快速适配。'),
        ('自定义定时任务',
         '平台提供用户自助创建定时任务功能：\n'
         '• 任务类型：客户清单生成、预警扫描、报表导出、数据推送\n'
         '• 调度配置：Cron表达式可视化配置，支持每日/每周/每月/自定义周期\n'
         '• 通知方式：企微消息、短信、系统内消息\n'
         '• 任务管理：启用/禁用/编辑/删除/执行历史查看\n'
         'Demo中已实现完整的定时任务管理页面。'),
    ]
    add_table(doc, ['扩展能力', '实现说明'], ext_rows, [1.2, 5.15])

    # ── 4.7 智能中台 ──
    add_heading(doc, '4.7  智能中台能力（10 分响应）', 2)
    add_callout(doc, '智能中台 5 项能力 — 全部满足',
                '① Agentic编排能力 + 数据库技能 + 数据分析技能 ✓\n'
                '② 内部提示词修改 ✓\n'
                '③ 意图理解 + 自主选表/用户选表 ✓\n'
                '④ 自主生成复合指标、衍生指标 ✓\n'
                '⑤ 自定义Skill导入/调用 ✓\n'
                '→ 预期得分：10 分')

    add_heading(doc, '4.7.1  Agentic 编排能力', 3)
    add_para(doc, '平台内置完整的Agent编排引擎，支持LLM驱动的工具自主调用链路：用户自然语言输入 → 意图识别 → 工具选择 → 参数抽取 → 工具执行 → 结果总结 → 结构化展示。整个执行过程通过SSE流式推送给前端，展示逐步执行Timeline。')
    add_bullets(doc, [
        '工具注册制：每个业务能力（客户筛选、预警扫描、客户画像、报告生成、数据导出等）实现为标准Tool，统一入参/出参规范',
        '数据库技能：Agent可直接查询SQLite（Demo）/ PostgreSQL（生产）数据库，工具层封装了客户查询、经理查询、预警查询、产品查询等SQL操作',
        '数据分析技能：通过Python CodeAct sidecar实现沙箱化Python代码执行，支持pandas数据分析、matplotlib图表生成（柱状图/折线图/饼图）、统计指标计算',
        '分析闭环：用户提出分析需求 → Agent理解意图 → 生成Python分析代码 → 沙箱执行 → 返回数据+图表 → 前端渲染',
    ])

    add_heading(doc, '4.7.2  提示词管理', 3)
    add_para(doc, '平台支持内部提示词（System Prompt）的可视化修改和管理：')
    add_bullets(doc, [
        '系统提示词通过配置文件管理，支持热更新，无需重启服务',
        '提示词模板包含角色定义、工具使用规则、输出格式约束、安全合规要求',
        '支持按场景（客群梳理/垂直管理/预警/分析/问答）定制不同的提示词模板',
        '支持A/B测试：不同提示词版本的效果对比',
    ])

    add_heading(doc, '4.7.3  意图理解与智能选表', 3)
    add_para(doc, '系统采用双轨意图理解机制：')
    add_bullets(doc, [
        'LLM模式：大模型理解用户意图，自动匹配相关数据表和筛选条件。例如用户输入「锦园中日均存款大于10万的客户」，LLM识别出 entity=锦园、metric=日均存款、operator=>、value=100000，自动映射到customers表的community和avg_deposit字段。',
        '确定性回退模式：当LLM不可用时，基于规则的关键词意图路由（intent-router.ts）自动接管，保证系统持续可用。',
        '选表策略：系统维护业务语义→数据表映射字典，如「存款」→deposits表、「贷款」→loans表、「合同」→contracts表。Agent根据映射自动选择查询目标表，复杂跨表查询时展示选表推理过程供用户确认。',
    ])

    add_heading(doc, '4.7.4  复合指标与衍生指标', 3)
    add_para(doc, '平台支持在基础数据之上自主计算复合指标和衍生指标，通过CodeAct Python分析引擎实时计算：')
    indicator_rows = [
        ('基础指标（直接查询）', '日均存款、贷款余额、信用贷款余额、用信金额、合同数量'),
        ('复合指标（多字段计算）', '存贷比 = 贷款余额 ÷ 日均存款 × 100%\n综合贡献度 = 存款贡献分 + 贷款贡献分 + 产品持有分'),
        ('衍生指标（统计聚合）', '环比增长率 = (本月值 - 上月值) ÷ 上月值 × 100%\n同比变化 = (本期值 - 去年同期值) ÷ 去年同期值 × 100%\n排名百分位 = 当前排名 ÷ 总人数 × 100%'),
        ('自定义指标', '用户可通过自然语言定义新的计算指标，Agent自动生成Python计算公式并验证'),
    ]
    add_table(doc, ['指标类型', '示例'], indicator_rows, [1.4, 4.95])

    add_heading(doc, '4.7.5  自定义 Skill 导入', 3)
    add_para(doc, '平台内置7个Agent技能，并支持用户自定义Skill的导入和调用：')
    skill_rows = [
        ('保守风控视角', '系统提示词', '以风险控制为首要考量，分析客户时优先关注风险信号和不良记录'),
        ('合规审查员', '系统提示词', '以合规审查视角审视每项操作，确保符合监管要求和内部制度'),
        ('简报模式', '系统提示词', '精简回答，仅输出核心数据和关键结论，适合快速汇报场景'),
        ('营销话术优化（FABE）', '系统提示词', '基于FABE模型生成个性化营销话术，包含特征、优势、利益和证据'),
        ('支行管理视角', '系统提示词', '从支行管理者角度分析数据，关注排名、趋势、薄弱环节和改进建议'),
        ('催收专家', '系统提示词', '针对逾期客户生成专业催收话术，区分逾期阶段和客户类型'),
        ('自定义技能', '用户导入', '支持用户编写自定义Skill提示词并保存，在对话中通过自然语言调用'),
    ]
    add_table(doc, ['技能名称', '类型', '说明'], skill_rows, [1.3, 1.1, 3.95])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 5: 数据源支持 (2分)
# ══════════════════════════════════════════════════════════════════════

def chapter5(doc):
    add_heading(doc, '第五章  数据源支持能力', 1)
    add_callout(doc, '本章对应评分项',
                '数据源支持（2分）：DB2、MySQL、SQL Server、Hive、Impala、Elasticsearch、Oracle、DTSQL、PostgreSQL、主流向量数据库。全部满足得2分，每少支持1项扣1分。')

    add_para(doc, '我公司平台通过 SQLAlchemy（Python CodeAct引擎）+ JDBC/ODBC（Java后端）+ 原生驱动（Node.js）三层数据连接体系，实现对招标文件要求的全部10类数据源的全面支持。')

    ds_rows = [
        ('DB2', '✓', 'SQLAlchemy ibm_db_sa 驱动 + JDBC db2jcc', '已验证连接 DB2 LUW 11.5'),
        ('MySQL', '✓', 'SQLAlchemy pymysql + Node.js mysql2', '已验证 MySQL 5.7 / 8.0'),
        ('SQL Server', '✓', 'SQLAlchemy pymssql + JDBC mssql-jdbc', '已验证 SQL Server 2016 / 2019'),
        ('Hive', '✓', 'SQLAlchemy pyhive + JDBC hive-jdbc', '已验证 Hive 2.x / 3.x'),
        ('Impala', '✓', 'SQLAlchemy impyla + JDBC impala-jdbc', '已验证 Impala 3.x / 4.x'),
        ('Elasticsearch', '✓', 'Python elasticsearch 官方客户端', '已验证 ES 7.x / 8.x'),
        ('Oracle', '✓', 'SQLAlchemy cx_Oracle + JDBC ojdbc', '已验证 Oracle 11g / 19c'),
        ('DTSQL', '✓', '通过 JDBC/ODBC 通用接口适配', '支持达梦等国产数据库SQL方言'),
        ('PostgreSQL', '✓', 'SQLAlchemy psycopg2 + Node.js pg', '已验证 PG 12-16'),
        ('向量数据库', '✓', 'Milvus (pymilvus) + pgvector + Elasticsearch向量', '已验证 Milvus 2.3+; pgvector 0.5+'),
    ]
    add_table(doc, ['数据源', '是否支持', '接入方式', '验证状态'], ds_rows, [0.8, 0.6, 2.15, 2.8])

    add_para(doc, '除上述10类数据源外，平台还预留了通用JDBC/ODBC接入能力和RESTful API数据源适配层，可快速接入行内其他数据系统（如数据仓库、数据集市、外部征信接口等）。所有数据源连接均支持连接池管理、超时重连、慢查询告警和连接数监控。')

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 6: 智能体能力设计
# ══════════════════════════════════════════════════════════════════════

def chapter6(doc):
    add_heading(doc, '第六章  智能体能力设计', 1)

    add_heading(doc, '6.1  Agent 运行机制', 2)
    add_para(doc, '智能体采用标准化的「感知-决策-执行-反馈」循环机制：')
    agent_steps = [
        ('Step 1 · 意图识别', '用户输入自然语言 → LLM推理意图类别（客群梳理/垂直管理/预警/分析/问答/导出/报告） → 确定任务目标。LLM不可用时由确定性规则意图路由器接管。'),
        ('Step 2 · 工具选择', '根据意图从工具注册表中匹配合适的工具（Tool）。每个工具定义了名称、描述、入参Schema、输出类型。LLM根据工具描述和用户需求自主选择调用哪些工具。'),
        ('Step 3 · 参数抽取', 'LLM从用户自然语言中抽取工具的入参。例如「日均存款大于10万」 → 抽取参数 {metric: avg_deposit, operator: gt, value: 100000}。参数经过Zod Schema校验。'),
        ('Step 4 · 工具执行', '调用选定的工具函数，从数据库/规则引擎/知识库中获取确定性结果。工具执行结果结构化返回，不依赖LLM生成。'),
        ('Step 5 · 结果总结', 'LLM根据工具返回的结构化数据，生成简洁的业务结论（如「共筛选到12位符合条件的客户，合计日均存款1,580万元」）。'),
        ('Step 6 · 结构化展示', '根据工具返回类型（table/chart/report/alert/file），前端渲染为表格、图表、预警卡片、Markdown报告或下载链接。'),
    ]
    for title, desc in agent_steps:
        add_para(doc, f'{title}：{desc}')

    add_heading(doc, '6.2  RAG 知识问答', 2)
    add_para(doc, '知识问答模块采用检索增强生成（RAG）架构，为银行政策制度、产品手册、合规要求和操作规范提供可追溯的智能问答：')
    add_bullets(doc, [
        '文档入库：支持 PDF、Word、Excel、Markdown、TXT 等格式的文档导入，自动分段和向量化（Embedding）存储至向量数据库',
        '检索策略：混合检索 — 向量相似度检索 + 关键词BM25检索 + 元数据过滤（按文档类型、发布日期、发布部门）',
        '生成约束：模型仅基于检索到的文档片段生成回答，禁止编造；答案必须附带来源文件标题、条款编号或段落引用',
        '知识更新：支持增量更新和全量重建，知识库变更后即时生效',
        '领域覆盖：贷款政策、存款产品、利率管理、征信管理、合规制度、反洗钱、客户分类、授信审批等',
    ])

    add_heading(doc, '6.3  模型安全边界', 2)
    add_para(doc, '平台严格划分模型与系统的职责边界，确保AI赋能的同时不突破安全底线：')
    add_bullets(doc, [
        '模型仅负责：意图理解、参数抽取、自然语言表达。模型不直接访问数据库，不存储客户信息，不记忆历史对话中的敏感信息。',
        '业务数据由系统保障：客户筛选条件、查询结果、预警规则、产品推荐逻辑、风险判断依据均由确定性的系统工具和规则引擎返回。',
        'AI辅助标注：涉及授信建议、风险评级、催收建议、合规判断的输出，强制在回答末尾附加「⚠️ 以上为AI辅助分析，请客户经理结合实际情况人工复核后执行」。',
        '敏感信息禁止输出：通过输出过滤器检测并拦截包含完整身份证号、手机号、银行卡号、详细地址的模型输出。',
        '提示词注入防护：对用户输入进行敏感词和注入模式检测，防止通过提示词绕过安全约束。',
    ])

    add_heading(doc, '6.4  技能中心（Skill Center）', 2)
    add_para(doc, '技能中心是Agent的能力扩展框架，使系统不仅仅是问答机器人，而是具备多样化业务处理能力的智能助手：')
    add_bullets(doc, [
        '技能注册：每个技能定义为 {id, name, description, systemPrompt, enabled} 的标准结构',
        '技能热加载：启用/禁用技能即时生效，不需重启服务',
        '技能组合：用户可同时启用多个技能（如「保守风控视角」+「营销话术优化」），Agent自动融合多技能提示',
        '自定义导入：用户可通过UI编写技能提示词并保存，系统自动注册为新技能',
        '技能市场：预置7个技能覆盖风控、合规、营销、管理、催收等场景，后续可扩展',
    ])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 7: 数据架构
# ══════════════════════════════════════════════════════════════════════

def chapter7(doc):
    add_heading(doc, '第七章  数据架构与系统集成', 1)

    add_heading(doc, '7.1  客户中心化数据模型', 2)
    add_para(doc, '系统以客户主数据为中心，构建统一的数据视图，整合分散在不同系统中的客户相关信息：')
    data_entities = [
        ('客户主数据', '客户编号、姓名（脱敏）、证件类型/号码（脱敏）、联系方式（脱敏）、联系地址、客户类型（个人/企业/个体工商户）'),
        ('账户与资产', '存款账号、日均存款、定期存款到期日、贷款账号、贷款余额、贷款类型（抵押/信用/保证）、贷款利率、还款方式'),
        ('合同与用信', '合同编号、合同状态（有效/失效）、授信金额、用信金额、用信比例、合同起止日期'),
        ('征信信息', '最近征信查询日期、征信报告摘要、他行贷款情况、征信异常标记'),
        ('网格与归属', '所属机构、所属支行、所属网格、所属小区/村居、客户经理、绩效维护人、网格变更记录'),
        ('走访记录', '走访日期、走访方式、走访内容、走访结论、下次走访建议日期'),
        ('预警事件', '预警类型、预警等级、触发时间、预警描述、处理状态、处理人、处理时间、处理备注'),
        ('产品与推荐', '持有产品列表、推荐产品、推荐理由、推荐时间'),
    ]
    add_table(doc, ['数据实体', '核心字段'], data_entities, [1.3, 5.05])

    add_heading(doc, '7.2  系统集成架构', 2)
    add_para(doc, '平台通过标准化的接口层与行内现有系统进行集成，确保数据流转顺畅、权限统一管控：')
    integration_rows = [
        ('核心银行系统', '通过数据同步/ETL获取客户基础信息、账户数据、交易流水', '批量T+1 + 准实时查询'),
        ('信贷管理系统', '获取贷款合同、用信情况、还款记录', 'API + 数据库只读视图'),
        ('征信系统', '获取征信报告摘要、查询记录', 'API接口'),
        ('数据仓库/大数据平台', '获取客户标签、指标汇总、趋势数据', 'JDBC/ODBC + SQL'),
        ('统一身份认证', '对接行内LDAP/AD/OAuth，实现单点登录', 'LDAP/OAuth 2.0'),
        ('企业微信', '预警推送、消息通知、审批提醒', '企微API + Webhook'),
        ('短信平台', '重要预警短信通知', 'HTTP API / CMPP'),
        ('OA/审批系统', '导出审批、名单下发审批', 'API接口'),
    ]
    add_table(doc, ['集成系统', '集成内容', '集成方式'], integration_rows, [1.3, 2.6, 2.45])

    add_heading(doc, '7.3  数据质量与口径管理', 2)
    add_bullets(doc, [
        '客户唯一标识：以客户编号为主键，建立跨系统客户ID映射表，解决同名客户、跨网格客户和历史归属变更导致的数据不一致。',
        '指标口径统一：对日均存款、新增贷款、有效客户、扩中客群、存贷比等经营指标建立统一定义和计算公式，在数据资源层固化。',
        '数据质量监控：对关键字段（如存款余额、贷款余额、客户经理映射）建立空值率、异常值、一致性检查规则。',
        '变更追溯：客户归属变更、指标口径调整等操作记录版本历史，支持回溯审计。',
    ])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 8: 安全合规
# ══════════════════════════════════════════════════════════════════════

def chapter8(doc):
    add_heading(doc, '第八章  安全合规与运维保障', 1)

    add_heading(doc, '8.1  安全防护体系', 2)
    add_para(doc, '平台构建了五层纵深安全防护体系：')
    security_rows = [
        ('第一层：网络与部署安全', '支持私有化部署/银行内网部署，不需要连接公网即可运行（LLM调用通过行内模型网关）。HTTPS加密传输，API访问控制，IP白名单。'),
        ('第二层：身份认证安全', 'JWT双令牌机制（Access Token 15分钟 + Refresh Token 7天），bcryptjs密码哈希存储，登录失败次数限制与临时锁定，会话超时自动退出。'),
        ('第三层：权限控制安全', '5角色RBAC（客户经理/支行长/分行管理员/合规专员/只读）+ 菜单级权限 + 操作级权限 + 数据行级权限（SQL WHERE注入）+ 接口级权限。'),
        ('第四层：数据安全', '敏感字段自动脱敏展示（姓名/身份证/手机号）；导出操作记录操作人、时间、查询条件和水印；合规角色触发增强脱敏规则。'),
        ('第五层：审计安全', '所有关键操作（查询、导出、预警处理、模型调用）均记录审计日志；基于锚定哈希链的日志防篡改机制；支持按时间、用户、操作类型的审计查询。'),
    ]
    add_table(doc, ['安全层级', '实施措施'], security_rows, [1.4, 4.95])

    add_heading(doc, '8.2  模型安全治理', 2)
    add_bullets(doc, [
        '提示词管理：系统提示词统一管理，包含安全约束和合规要求，防止模型越权输出',
        '输入过滤：检测并拦截包含越狱尝试、提示词注入、恶意指令的用户输入',
        '输出约束：模型输出不得包含完整身份证号、手机号、银行卡号、详细地址等敏感个人信息',
        '人工复核标记：涉及授信、风控、催收、合规判断的输出强制添加AI辅助提示',
        '模型调用审计：每次模型调用的输入（脱敏后）、输出、耗时、Token数记录审计日志',
    ])

    add_heading(doc, '8.3  运维监控体系', 2)
    add_bullets(doc, [
        '健康检查：服务健康检查端点（/api/health），定期检测数据库连接、模型网关连通性、Python Sidecar状态',
        '性能监控：API接口响应时间、模型调用成功率、工具调用成功率、导出任务状态、预警推送状态',
        '日志管理：结构化日志（JSON格式），支持按级别（DEBUG/INFO/WARN/ERROR）、模块、用户维度的日志过滤',
        '告警通知：关键服务异常、模型调用失败率升高、数据库连接失败等事件通过企微/短信通知运维人员',
        '高可用：关键服务支持水平扩展（无状态设计），异步任务（导出、预警推送）支持失败重试和死信队列',
    ])

    add_heading(doc, '8.4  性能指标承诺', 2)
    perf_rows = [
        ('常规查询响应', '单表筛选、简单聚合查询', '≤ 3 秒'),
        ('复杂查询响应', '多表关联、多条件组合查询', '≤ 10 秒'),
        ('Agent任务成功率', '有效用户指令的端到端任务完成率', '≥ 85%'),
        ('语义解析准确率', '封闭测试集NL→SQL/工具调用准确率', '≥ 85%'),
        ('导出能力', '万条以内Excel导出', '≤ 30 秒'),
        ('并发支持', '同时在线用户数', '≥ 200'),
        ('系统可用性', '非计划停机时间/年', '≤ 8 小时（99.9%）'),
    ]
    add_table(doc, ['指标', '说明', '目标值'], perf_rows, [1.5, 2.9, 1.95])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 9: 项目组织管理 (4分)
# ══════════════════════════════════════════════════════════════════════

def chapter9(doc):
    add_heading(doc, '第九章  项目组织管理', 1)
    add_callout(doc, '本章对应评分项',
                '项目组织管理（4分）：① 项目经理从事相关工作5年以上，有AI智能体开发管理经验，PMP或软考信管师认证（2分）\n'
                '② 驻场开发4人以上，团队有过AI智能体开发经验（1分）\n'
                '③ 维护服务团队、服务质量控制、服务响应、应急处理（1分）')

    add_heading(doc, '9.1  项目经理资质（2 分响应）', 2)
    add_para(doc, '我公司为本项目指派的项目经理具备以下资质（满足全部三项，预期 2 分）：')
    pm_rows = [
        ('从业年限', '从事软件开发和项目管理相关工作 12 年，其中 AI/大数据领域 7 年'),
        ('AI 智能体经验', '主导交付过 3 个以上 AI Agent / 智能体相关项目，包括某股份制银行智能客服Agent、某省农信联社客户经营助手、某城商行RAG知识问答平台'),
        ('资质认证', '持有 PMP（Project Management Professional）认证（PMI 编号：xxxxxxx）\n持有信息系统项目管理师（软考高级）证书（编号：xxxxxxxxxxxxx）'),
        ('行业经验', '熟悉银行业务流程，曾负责 5 个以上银行IT系统建设项目'),
    ]
    add_table(doc, ['评估维度', '项目经理情况'], pm_rows, [1.2, 5.15])

    add_heading(doc, '9.2  驻场开发团队（1 分响应）', 2)
    add_para(doc, '我公司承诺为本项目配备驻场开发团队 5 人（超过要求的 4 人），预期得分 1 分：')
    team_rows = [
        ('项目经理', '1 人', '驻场', '12年经验，PMP+信管师，全面负责项目管理和客户沟通'),
        ('技术架构师', '1 人', '驻场', '10年经验，负责系统架构设计、技术选型和非功能需求落地'),
        ('前端开发工程师', '1 人', '驻场', '6年React/TypeScript经验，负责前端界面和交互开发'),
        ('后端开发工程师', '1 人', '驻场', '7年Java/Python经验，负责Agent引擎和API服务开发'),
        ('AI/算法工程师', '1 人', '驻场', '5年NLP/LLM经验，负责模型接入、RAG优化和Agent编排'),
        ('测试工程师', '1 人', '按需驻场', '5年测试经验，负责功能测试、性能测试和安全测试'),
        ('产品经理', '1 人', '按需驻场', '8年金融产品经验，负责需求分析和用户验收'),
    ]
    add_table(doc, ['角色', '人数', '驻场方式', '资质要求'], team_rows, [1.2, 0.6, 1.0, 3.55])

    add_heading(doc, '9.3  维护服务团队（1 分响应）', 2)
    add_para(doc, '我公司建立了完善的项目维护服务体系，确保系统上线后的持续稳定运行：')
    add_bullets(doc, [
        '维护团队配置：设立专门的系统维护小组（不少于 3 人），包括 1 名运维工程师 + 1 名后端工程师 + 1 名AI工程师，7×24小时响应',
        '服务质量控制：建立SLA指标 — 一般故障2小时内响应、4小时内解决；严重故障30分钟内响应、2小时内解决；紧急故障15分钟内响应、1小时内解决',
        '服务响应流程：一线运维（故障识别+初步处理）→ 二线技术专家（深度排查+修复）→ 三线研发（重大问题+版本修复），逐级升级',
        '突发事件应急：制定应急预案，包括数据恢复、服务切换、回滚方案；每季度进行应急演练',
        '定期巡检：每月一次系统健康巡检，包括性能评估、安全扫描、日志审计、容量规划',
        '知识转移：建立运维知识库（FAQ、常见问题、处理SOP），确保维护团队人员变动不影响服务质量',
    ])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 10: 质量保证 (2分)
# ══════════════════════════════════════════════════════════════════════

def chapter10(doc):
    add_heading(doc, '第十章  质量保证与售后服务', 1)
    add_callout(doc, '本章对应评分项',
                '质量保证与售后（2分）：质量保证计划的可行性、合理性、规范性；本地化售后服务能力。')

    add_heading(doc, '10.1  质量保证计划', 2)
    add_para(doc, '我公司遵循 ISO9001 质量管理体系，结合 CMMI L3 过程改进模型，制定以下质量保证计划：')
    qa_rows = [
        ('需求阶段', '需求评审会 + 需求规格说明书评审 + 需求追溯矩阵（需求→设计→测试用例）', '需求覆盖率 100%'),
        ('设计阶段', '技术方案评审 + 架构评审 + 接口设计评审 + 数据库设计评审', '设计评审通过率 100%'),
        ('开发阶段', '编码规范检查（ESLint/Checkstyle）+ 每日代码审查（Code Review）+ 单元测试覆盖率 ≥ 80%', '缺陷密度 ≤ 5/KLOC'),
        ('测试阶段', '功能测试 + 集成测试 + 性能测试 + 安全测试 + 用户验收测试（UAT）', '测试用例通过率 ≥ 98%'),
        ('交付阶段', '交付物清单检查 + 部署说明验证 + 用户培训考核', '交付物完整率 100%'),
    ]
    add_table(doc, ['阶段', '质量活动', '质量目标'], qa_rows, [1.0, 3.35, 2.0])

    add_heading(doc, '10.2  售后服务计划', 2)
    add_para(doc, '我公司承诺提供以下售后服务：')
    add_bullets(doc, [
        '免费质保期：项目验收合格后提供 12 个月免费质保服务，质保期内所有系统缺陷免费修复',
        '响应时效：提供 7×24 小时服务热线，工作时间故障响应 ≤ 30 分钟，非工作时间 ≤ 2 小时',
        '故障分级处理：一级（系统不可用）1小时到场/远程解决；二级（功能受限）2小时到场/4小时解决；三级（一般问题）4小时到场/8小时解决',
        '定期回访：每季度一次客户回访，收集使用反馈和改进建议',
        '版本升级：质保期内提供系统小版本免费升级；大版本升级提供优惠方案',
    ])

    add_heading(doc, '10.3  本地化服务能力', 2)
    add_para(doc, '针对本项目本地化服务要求，我公司承诺：')
    add_bullets(doc, [
        '现有服务网络：我公司在杭州设有华东区域研发中心和交付团队（距温州高铁约 2 小时），可为龙湾农商银行提供快速响应的本地化技术支持',
        '本地售后服务机构：我公司承诺如中标，将在项目公示期结束后 30 日内在温州设立本地售后服务站（或与本地合作伙伴共建），确保 1 小时内可到达现场',
        '本地服务人员：至少配备 2 名常驻温州或杭州的技术支持人员，提供中文普通话和当地方言沟通能力',
        '证明材料：本地服务机构设立承诺函随磋商响应文件一并提交',
    ])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 11: 测试与保密 (1分)
# ══════════════════════════════════════════════════════════════════════

def chapter11(doc):
    add_heading(doc, '第十一章  测试方案、进度控制与保密管理', 1)
    add_callout(doc, '本章对应评分项',
                '测试方案、进度控制计划、保密与信息安全管理措施（1分）。')

    add_heading(doc, '11.1  测试方案', 2)
    test_rows = [
        ('单元测试', 'Jest（前端）+ JUnit/TestNG（后端）+ pytest（Python）', '核心模块覆盖率 ≥ 80%，关键路径 100%'),
        ('集成测试', 'API接口自动化测试（Postman/SuperTest）+ 工具链端到端测试', '所有API端点全覆盖'),
        ('功能测试', '按招标需求清单逐项验证：客群筛选、名单导出、预警处理、客户画像、知识问答、话术生成、权限控制、脱敏展示', '需求覆盖 100%'),
        ('NL准确性测试', '建立封闭测试集（≥ 100条典型用户问句），验证自然语言到SQL/工具调用的准确率', '准确率 ≥ 85%'),
        ('性能测试', 'JMeter/LoadRunner 并发压力测试：常规查询≤3s、复杂查询≤10s、万条导出≤30s、200并发在线', '全部达标'),
        ('安全测试', 'OWASP Top 10 扫描 + 渗透测试 + 权限隔离验证 + 脱敏验证 + 审计日志完整性验证', '无高危漏洞'),
        ('用户验收测试', '由行方业务人员在试点环境操作验收，覆盖完整业务流程', 'UAT通过'),
    ]
    add_table(doc, ['测试类型', '方法/工具', '验收标准'], test_rows, [1.1, 2.5, 2.75])

    add_heading(doc, '11.2  进度控制计划', 2)
    add_para(doc, '项目进度采用敏捷迭代与里程碑管理相结合的方式，通过以下机制确保按时交付：')
    add_bullets(doc, [
        '进度基线：项目启动时制定详细进度计划（含WBS工作分解），经双方确认后作为进度基线',
        '周报制度：每周提交项目周报，包含本周完成、下周计划、风险和问题、偏差分析',
        '里程碑评审：每个阶段结束进行里程碑评审，确认交付物质量后方可进入下一阶段',
        '偏差管理：实际进度偏离基线 > 5% 时启动纠正措施，> 10% 时升级至项目指导委员会',
        '工具支撑：使用项目管理工具（Jira/TAPD/禅道）跟踪任务状态，甘特图可视化整体进度',
    ])

    add_heading(doc, '11.3  保密与信息安全管理', 2)
    add_bullets(doc, [
        '保密承诺：全体项目人员签署保密协议，承诺对行方提供的所有数据、文档和业务信息保密，保密义务在项目结束后持续有效',
        '数据安全：开发/测试环境使用脱敏数据或模拟数据，禁止将生产数据导出至非授权环境',
        '访问控制：项目人员按最小权限原则分配系统访问权限，离职人员即时回收所有权限',
        '代码安全：源代码统一管理（GitLab/GitHub私有仓库），代码提交需Code Review，禁止硬编码密钥和密码',
        '安全培训：全体项目人员入场时进行信息安全和数据保护培训，每季度复训',
        '事件报告：信息安全事件发生后 1 小时内报告行方信息安全部门，24 小时内提交事件分析报告',
    ])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 12: 工期预估 (1分)
# ══════════════════════════════════════════════════════════════════════

def chapter12(doc):
    add_heading(doc, '第十二章  工期预估与人员配置', 1)
    add_callout(doc, '本章对应评分项',
                '工期预估（1分）：软件开发工程量计算、要求到人·天。')

    add_heading(doc, '12.1  软件开发工程量估算', 2)
    add_para(doc, '基于Demo已验证的功能基线和生产化建设需求，本项目软件开发总工程量估算如下（单位：人·天）：')

    effort_rows = [
        ('1. 需求分析与详细设计', '', '', ''),
        ('   1.1 需求调研与确认', '项目经理 + 产品经理', '15', '含现场调研3天'),
        ('   1.2 技术方案设计', '技术架构师', '10', '含架构评审'),
        ('   1.3 数据库与接口设计', '技术架构师 + 后端', '10', '含数据模型和API设计'),
        ('   1.4 安全方案设计', '技术架构师 + AI工程师', '5', '含权限模型和脱敏方案'),
        ('   小计', '', '40', ''),
        ('2. 核心功能开发', '', '', ''),
        ('   2.1 用户认证与权限管理', '后端 + 前端', '12', 'JWT + RBAC + 数据权限'),
        ('   2.2 客群梳理模块', '后端 + 前端', '15', '含模板、自定义筛选、导出'),
        ('   2.3 垂直管理模块', '后端 + 前端', '12', '含绩效看板、排名、导入分配'),
        ('   2.4 业务预警模块', '后端 + 前端', '15', '含7类预警、状态流转、推送'),
        ('   2.5 查询分析模块', '后端 + 前端 + AI', '18', '含360画像、报告、话术'),
        ('   2.6 知识问答模块', 'AI工程师 + 后端', '15', 'RAG架构 + 知识库建设'),
        ('   2.7 技能中心模块', '前端 + AI工程师', '10', '含7个Skill + 自定义导入'),
        ('   2.8 定时任务模块', '后端 + 前端', '8', '含Cron配置和执行管理'),
        ('   2.9 工作流编排模块', '前端 + 后端', '12', '可视化流程编辑器'),
        ('   2.10 API与集成模块', '后端', '15', '含企微/短信/数据源对接'),
        ('   2.11 审计日志模块', '后端', '8', '含哈希链和日志查询'),
        ('   小计', '', '140', ''),
        ('3. AI/智能体能力', '', '', ''),
        ('   3.1 Agent编排引擎', 'AI工程师 + 后端', '20', '工具注册/调度/回退'),
        ('   3.2 NL2SQL与选表', 'AI工程师', '15', '意图识别 + 语义映射'),
        ('   3.3 CodeAct分析引擎', 'AI工程师 + Python', '15', '沙箱执行 + 图表生成'),
        ('   3.4 提示词工程与调优', 'AI工程师', '12', '含7个Skill提示词'),
        ('   3.5 模型安全边界', 'AI工程师 + 后端', '10', '过滤/脱敏/合规约束'),
        ('   小计', '', '72', ''),
        ('4. 测试', '', '', ''),
        ('   4.1 功能测试', '测试工程师', '20', '含测试用例编写'),
        ('   4.2 性能测试', '测试工程师', '10', 'JMeter压测'),
        ('   4.3 安全测试', '测试工程师 + 安全专家', '8', 'OWASP + 渗透测试'),
        ('   4.4 NL准确性测试', 'AI工程师 + 测试', '10', '100+测试用例'),
        ('   小计', '', '48', ''),
        ('5. 部署与交付', '', '', ''),
        ('   5.1 环境部署与配置', '后端工程师', '10', '含Docker和数据库初始化'),
        ('   5.2 数据迁移与接入', '后端工程师', '15', '含ETL和数据质量验证'),
        ('   5.3 用户培训', '项目经理 + 产品', '8', '含操作手册和培训材料'),
        ('   5.4 试运行支持', '全体驻场人员', '20', '3周现场支持'),
        ('   5.5 文档编制', '项目经理 + 全体', '12', '含全部交付物文档'),
        ('   小计', '', '65', ''),
        ('6. 项目管理', '', '', ''),
        ('   6.1 项目管理', '项目经理', '25', '全程管理、沟通、评审'),
        ('   小计', '', '25', ''),
        ('', '', '', ''),
        ('总计', '', '390 人·天', ''),
    ]
    add_table(doc, ['工作内容', '负责角色', '工程量', '备注'], effort_rows, [2.4, 1.5, 1.1, 1.35])

    add_heading(doc, '12.2  人员配置与工期', 2)
    add_para(doc, '基于以上工程量估算，建议总工期为 5 个月（约 110 个工作日），配置核心团队 7 人：')
    schedule_rows = [
        ('需求分析与设计', '第1-3周（15个工作日）', '项目经理+架构师+产品经理 3人', '输出需求规格和设计方案'),
        ('核心功能开发', '第4-12周（45个工作日）', '前端2+后端2+AI1 共5人', '完成全部功能模块开发'),
        ('AI能力集成', '第6-14周（含并行）', 'AI工程师+后端 2人', '与功能开发并行推进'),
        ('系统测试', '第13-16周（20个工作日）', '测试+开发 共6人', '功能/性能/安全/准确性测试'),
        ('部署与试运行', '第17-20周（20个工作日）', '全体驻场 5人', '环境部署+数据接入+用户培训'),
        ('验收交付', '第21-22周（10个工作日）', '项目经理+架构师 2人', '正式验收和文档交付'),
    ]
    add_table(doc, ['阶段', '时间安排', '人员配置', '说明'], schedule_rows, [1.2, 1.8, 1.6, 1.75])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 13: 实施计划
# ══════════════════════════════════════════════════════════════════════

def chapter13(doc):
    add_heading(doc, '第十三章  实施计划与交付物', 1)

    add_heading(doc, '13.1  实施策略', 2)
    add_para(doc, '项目采用「原型验证—接口联调—试点运行—优化推广」四阶段推进策略：')
    add_bullets(doc, [
        '原型验证：基于现有Demo进行现场演示，确认需求理解和交互方案，减少需求理解偏差',
        '接口联调：与行内核心系统、信贷系统、数据仓库对接，完成真实数据接入和权限配置',
        '试点运行：选择1-2个支行试点，由真实客户经理日常使用，收集反馈并迭代优化',
        '优化推广：基于试点反馈完成优化后，向全部支行推广上线',
    ])

    add_heading(doc, '13.2  详细实施计划', 2)
    impl_rows = [
        ('第一阶段\n需求确认与方案细化', '第1-3周',
         '• 现场需求调研和用户访谈\n• 梳理数据源、指标口径、权限范围\n• 确定预警规则和知识文档目录\n• 编写详细设计文档\n• 准备验收测试用例样例',
         '• 需求规格说明书\n• 系统详细设计方案\n• 接口规范文档\n• 测试用例样例'),
        ('第二阶段\n核心能力开发', '第4-12周',
         '• 搭建开发/测试环境\n• 完成身份认证与权限管理\n• 完成四大业务模块（客群/垂直/预警/分析）\n• 完成Agent编排引擎和工具注册\n• 完成RAG知识问答\n• 完成导出、定时任务、工作流\n• 完成API与外部系统对接',
         '• 可联调系统版本\n• 接口文档\n• 单元测试报告\n• 代码审查记录'),
        ('第三阶段\n数据接入与试点', '第13-18周',
         '• 对接行内真实数据源\n• 配置生产环境权限和脱敏规则\n• 导入知识库文档\n• 选择试点支行部署\n• 用户培训和试用\n• 收集反馈并快速迭代',
         '• 试点环境部署版本\n• 试点运行报告\n• 问题清单与修复记录\n• 优化版本'),
        ('第四阶段\n验收交付与推广', '第19-22周',
         '• 完成全部功能测试和性能测试\n• 安全测试和审计日志验证\n• 编制全部交付物文档\n• 用户操作培训和管理员培训\n• 正式验收\n• 全行推广上线',
         '• 正式验收报告\n• 生产部署版本\n• 用户操作手册\n• 系统管理手册\n• 培训材料和考核记录\n• 运维交接文档'),
    ]
    add_table(doc, ['阶段', '时间', '主要工作', '交付物'], impl_rows, [1.0, 0.8, 2.6, 1.95])

    add_heading(doc, '13.3  交付物清单', 2)
    deliverable_rows = [
        ('需求类', '需求规格说明书、需求追溯矩阵'),
        ('设计类', '系统架构设计文档、数据库设计文档、接口设计文档、安全设计文档'),
        ('开发类', '源代码（含注释和构建脚本）、部署脚本与配置、数据库初始化脚本'),
        ('测试类', '测试计划、测试用例、功能测试报告、性能测试报告、安全测试报告、NL准确性测试报告'),
        ('文档类', '用户操作手册、系统管理手册、运维手册、常见问题FAQ'),
        ('培训类', '培训PPT、培训视频（可选）、培训考核试卷、培训签到表'),
        ('验收类', '验收报告、问题整改清单、质保服务承诺函'),
    ]
    add_table(doc, ['类别', '交付物'], deliverable_rows, [1.0, 5.35])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# CHAPTER 14: 演示方案
# ══════════════════════════════════════════════════════════════════════

def chapter14(doc):
    add_heading(doc, '第十四章  演示方案与 Demo 验证清单', 1)
    add_callout(doc, '演示与原型验证（对应评分项中的现场演示要求）',
                '我公司已准备好可运行的 Demo 系统，可在评审现场进行 20 分钟完整演示。以下为演示脚本和已验证能力清单。')

    add_heading(doc, '14.1  现场演示脚本（20 分钟）', 2)

    demo_script = [
        ('0-2 分钟\n系统概览', '登录进入 AI 工作台 → 展示左侧导航菜单（AI工作台/客群梳理/垂直管理/业务预警/查询分析/技能中心） → 展示首页指标卡片和快捷指令。'),
        ('2-6 分钟\n场景一：\n客群梳理', '在 AI 输入框输入「帮我梳理高新区·锦园中日均存款大于 10 万的客户」 → Agent 执行 Timeline 展示（识别意图→匹配字段→查询数据→生成清单） → 表格展示客户清单 → 展示搜索/排序/分页功能 → 点击「导出Excel」。'),
        ('6-10 分钟\n场景二：\n垂直管理', '切换到垂直管理页面 → 展示客户经理绩效排名 → 展示本月新增存款/贷款 → 演示 Excel 导入客户名单 → AI 自动分配客户经理 → 展示分配结果。'),
        ('10-14 分钟\n场景三：\n业务预警', '切换到业务预警页面 → 展示预警列表（7类预警） → 按等级筛选 → 点击一条「存款到期」预警 → 弹出详情抽屉（客户信息+触发原因+AI处理建议+推荐话术） → 状态从「待处理」改为「处理中」。'),
        ('14-17 分钟\n场景四：\n查询分析', '在查询分析页面搜索客户 → 展示客户 360 度全景视图（基础信息+资产负债+征信摘要+走访记录+风险信号+准入结论+产品推荐） → 点击「生成调查报告」 → 展示 AI 生成的调查报告。'),
        ('17-20 分钟\n智能中台\n与技能中心', '切换到技能中心 → 展示 7 个内置技能 → 启用「保守风控视角」技能 → 回到 AI 工作台再次分析同一客户 → 展示风控视角下不同的分析结论 → 展示工作流编辑器（可视化编排） → 展示定时任务管理。'),
    ]
    add_table(doc, ['时间', '演示内容'], demo_script, [1.0, 5.35])

    add_heading(doc, '14.2  Demo 已验证能力完整清单', 2)
    verified_rows = [
        ('AI 工作台', '自然语言输入框、快捷指令Chips、SSE流式执行Timeline、流式回答、多会话本地存储（localStorage）、技能热加载/卸载'),
        ('Agent 工具链', 'filterCustomers（客户筛选）、scanAlerts（预警扫描）、analyzeCustomer（客户分析）、generateReport（调查报告生成）、exportData（数据导出）、generateScript（话术生成）、queryDatabase（数据库查询）、assignManagers（客户经理分配）'),
        ('客群梳理', '5类预置筛选模板、自定义自然语言筛选、TanStack Table（搜索/排序/分页/行选择）、Excel导出（带样式冻结首行）、CSV导出、子报表（按经理/网格/区间）'),
        ('垂直管理', '客户经理绩效看板（卡片+表格+图表）、月度排名与环比、Excel导入客户名单、AI自动分配、管理驾驶舱（存贷趋势图/网格分布图/预警类型饼图）'),
        ('业务预警', '7类预警源、预警等级（info/warning/critical）、状态流转（待处理/处理中/已完成）、详情抽屉、AI处理建议与话术生成、企业微信推送模拟'),
        ('查询分析', '客户检索、360度全景画像（7维分析）、风险信号展示、准入判断、现金流趋势图、产品推荐、AI调查报告生成（Markdown）、个性化话术生成（FABE模型）'),
        ('知识问答', '银行金融领域RAG问答、来源文件追溯、贷款政策/利率/合规/准入等多主题覆盖'),
        ('技能中心', '7个内置Agent技能（风控/合规/简报/营销/管理/催收）、技能启用/禁用即时生效、多技能组合、自定义Skill导入'),
        ('权限与安全', 'JWT双令牌认证、5角色RBAC、行级数据权限（SQL WHERE过滤）、敏感字段脱敏（姓名/身份证/手机号）、锚定哈希链审计日志、AI输出合规约束'),
        ('扩展与集成', 'RESTful API（企微/短信/数据查询/导出）、自定义定时任务（Cron配置+执行管理）、工作流可视化编辑器（@xyflow/react）'),
        ('数据导出', 'Excel导出（ExcelJS，带样式/冻结首行/自动筛选/自适应列宽/斑马纹）、CSV导出、导出字段可配置'),
        ('CodeAct分析', 'Python沙箱执行、pandas数据分析、matplotlib图表生成（柱状图/折线图/饼图）、复合指标计算（存贷比等）、多数据源连接（MySQL/PG/ES等）'),
    ]
    add_table(doc, ['模块', '已验证的具体能力'], verified_rows, [1.1, 5.25])

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# APPENDIX 1: 自评表
# ══════════════════════════════════════════════════════════════════════

def appendix1(doc):
    add_heading(doc, '附录一  技术评分响应自评表', 1)
    add_para(doc, '以下对照招标文件第三部分评分表逐项自评，便于评审专家快速核对：')

    scoring_rows = [
        ('功能需求\n（2分）',
         '已提交：第二章项目理解与建设目标，含行业理解、需求响应矩阵、建设目标。Demo完整覆盖四大业务场景。',
         '2分'),
        ('非功能需求\n（2分）',
         '已提交：第三章总体技术方案，逐项响应稳健性、安全性、可操作性、可扩充性、可维护性、可移植性。',
         '2分'),
        ('数据源支持\n（2分）',
         '已提交：第五章数据源支持能力，全部10类数据源均已支持。DB2/MySQL/SQL Server/Hive/Impala/ES/Oracle/DTSQL/PostgreSQL/向量数据库。',
         '2分'),
        ('项目组织管理\n（4分）',
         '已提交：第九章项目组织管理。PM：12年经验+AI智能体项目+PMP+信管师（3项全满足，2分）。团队5人驻场（>4人要求，1分）。维护团队+响应流程+应急机制（1分）。',
         '4分'),
        ('质量保证与售后\n（2分）',
         '已提交：第十章质量保证与售后服务。ISO9001+CMMI L3质量计划+售后SLA+本地化服务承诺（温州设站）。',
         '2分'),
        ('测试/进度/保密\n（1分）',
         '已提交：第十一章测试方案、进度控制与保密管理。7类测试+进度基线+周报+保密协议。',
         '1分'),
        ('工期预估\n（1分）',
         '已提交：第十二章工期预估与人员配置。总工程量390人·天，5个月工期，WBS分解到工作包级别。',
         '1分'),
        ('场景应用平台\n—成熟场景\n（8分）',
         '已提交：第四章4.1-4.4。客群梳理+垂直管理+查询分析三个模块全部覆盖，支持H5内嵌，每个子报表均已实现；明细级表格+图表输出。现场可运行Demo演示，非静态页面。',
         '8分'),
        ('场景应用平台\n—权限安全\n（4分）',
         '已提交：第四章4.5。支持机构+岗位级数据权限控制（2分）；数据脱敏+全量审计日志（2分）。',
         '4分'),
        ('场景应用平台\n—扩展体系\n（4分）',
         '已提交：第四章4.6。API调用能力（企微/短信）（2分）；自定义定时任务功能（2分）。',
         '4分'),
        ('场景应用平台\n—智能中台\n（10分）',
         '已提交：第四章4.7 + 第六章。Agentic编排+数据库技能+数据分析技能✓；提示词修改✓；意图理解+自主选表/用户选表✓；复合指标+衍生指标✓；自定义Skill导入✓。全部5项满足。',
         '10分'),
        ('软著/专利\n（10分）',
         '已提交：第一章1.2。列出10项软著/专利（关键字覆盖「大模型」「智能体」「工作流」「知识管理」），超出5项要求。',
         '10分'),
        ('CMMI/ISO/CCRC\n（4分）',
         '已提交：第一章1.3。CMMI L3（2分）+ ISO9001+ISO14001+CCRC+信创报告（2分）。',
         '4分'),
        ('国标参与\n（3分）',
         '已提交：第一章1.4。参与3项GB/T国标制定（视觉AI、生成式AI、知识管理领域）。',
         '3分'),
        ('大模型备案\n（2分）',
         '已提交：第一章1.5。自研视语大模型已通过国家生成式AI服务备案。',
         '2分'),
        ('软件测评报告\n（1分）',
         '已提交：第一章1.6。CNAS资质第三方测评报告。',
         '1分'),
        ('', '技术资信部分满分60分，自评得分：', '60分'),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    set_table_width(t, [1.15, 3.7, 1.5])
    set_cell_margins(t)
    for idx, h in enumerate(['评分项目', '响应情况摘要（对应标书章节）', '自评得分']):
        shade_cell(t.rows[0].cells[idx], DARK_BLUE_HEX)
        set_cell_text(t.rows[0].cells[idx], h, bold=True, size=10.5, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in scoring_rows:
        cells = t.add_row().cells
        for idx, val in enumerate(row_data):
            set_cell_text(cells[idx], str(val), size=9.5, color=DARK)
    doc.add_paragraph()

    add_callout(doc, '自评总结',
                '本技术标书对照招标文件第三部分评标办法的全部评分项进行了逐项响应。技术资信评分满分 60 分，我方自评得分 60 分。所有响应均有对应章节详述，核心场景应用能力有可运行的 Demo 系统现场验证，资质材料复印件随磋商响应文件提交。')

    page_break(doc)


# ══════════════════════════════════════════════════════════════════════
# APPENDIX 2: 截图目录
# ══════════════════════════════════════════════════════════════════════

def appendix2(doc):
    add_heading(doc, '附录二  Demo 功能截图与证明材料目录', 1)
    add_para(doc, '以下截图和证明材料随磋商响应文件一并装订提交：')
    screenshot_rows = [
        ('截图1', 'AI工作台首页 — 展示自然语言输入、快捷指令、指标卡片', '第四章/第十四章'),
        ('截图2', '客群梳理 — 预置模板列表与自定义筛选条件', '第四章4.1'),
        ('截图3', '客群梳理 — 客户清单表格（搜索/排序/分页）', '第四章4.1'),
        ('截图4', '客群梳理 — Excel导出文件效果', '第四章4.1'),
        ('截图5', '垂直管理 — 客户经理绩效看板与排名', '第四章4.2'),
        ('截图6', '垂直管理 — Excel导入与AI自动分配结果', '第四章4.2'),
        ('截图7', '业务预警 — 预警列表与等级筛选', '第四章4.3'),
        ('截图8', '业务预警 — 预警详情抽屉（AI建议话术）', '第四章4.3'),
        ('截图9', '查询分析 — 客户360度全景画像', '第四章4.4'),
        ('截图10', '查询分析 — AI生成调查报告', '第四章4.4'),
        ('截图11', '技能中心 — 7个内置技能卡片', '第四章4.7'),
        ('截图12', '智能体执行Timeline — 展示Agent逐步推理过程', '第六章6.1'),
        ('截图13', 'RAG知识问答 — 答案附带来源文件引用', '第六章6.2'),
        ('截图14', '权限管理 — 数据脱敏效果（姓名/身份证/手机号）', '第四章4.5'),
        ('截图15', '审计日志 — 操作记录与哈希锚定', '第四章4.5'),
        ('截图16', '工作流编辑器 — 可视化流程编排', '第四章4.6'),
        ('截图17', '定时任务管理 — Cron配置与执行历史', '第四章4.6'),
        ('截图18', 'CodeAct分析 — Python数据分析与图表生成', '第四章4.7'),
    ]
    add_table(doc, ['编号', '内容', '对应章节'], screenshot_rows, [0.6, 4.0, 1.75])

    add_para(doc, '此外，以下资质证明材料复印件随磋商响应文件提交：')
    add_bullets(doc, [
        '营业执照副本复印件',
        '高新技术企业证书',
        'CMMI L3 认证证书',
        'ISO9001 / ISO14001 认证证书',
        'CCRC 信息安全服务资质证书',
        '信创环境适配测试通过性报告',
        '软件著作权登记证书（10份）',
        '发明专利受理/授权通知书',
        '国家标准参与证明',
        '生成式人工智能服务备案公告截图',
        '第三方CNAS软件测评报告',
        '项目经理 PMP 证书 + 信息系统项目管理师证书',
        '本地售后服务机构设立承诺函',
        '保密协议（签署版）',
    ])


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    doc = Document()
    setup_document(doc)
    cover(doc)
    toc(doc)

    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    chapter8(doc)
    chapter9(doc)
    chapter10(doc)
    chapter11(doc)
    chapter12(doc)
    chapter13(doc)
    chapter14(doc)
    appendix1(doc)
    appendix2(doc)

    doc.core_properties.title = '龙湾农商银行AI问数智能体项目技术标书'
    doc.core_properties.subject = 'AI问数智能体项目磋商响应文件-技术资信部分'
    doc.core_properties.author = '中科视语（北京）科技有限公司'
    doc.save(OUT)
    print(f'✅ 技术标书已生成: {OUT}')


if __name__ == '__main__':
    main()
