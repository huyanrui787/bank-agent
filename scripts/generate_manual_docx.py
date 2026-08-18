#!/usr/bin/env python3
"""Generate the AI customer assistant Word operation manual."""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
IMG_DIR = BASE_DIR / "generated" / "manual_screenshots"
OUT_FILE = BASE_DIR / "generated" / "AI客户经营助手_操作手册.docx"

SCREENSHOTS = {
    "login": IMG_DIR / "01_login.png",
    "dashboard_empty": IMG_DIR / "02_dashboard_empty.png",
    "dashboard_customer_segment": IMG_DIR / "03_dashboard_customer_segment.png",
    "dashboard_business_alert": IMG_DIR / "04_dashboard_business_alert.png",
    "dashboard_customer_analysis": IMG_DIR / "05_dashboard_customer_analysis.png",
    "dashboard_vertical_management": IMG_DIR / "06_dashboard_vertical_management.png",
    "customer_segments_top": IMG_DIR / "07_customer_segments_top.png",
    "customer_segments_list": IMG_DIR / "08_customer_segments_list.png",
    "vertical_management_top": IMG_DIR / "09_vertical_management_top.png",
    "vertical_management_assignment": IMG_DIR / "10_vertical_management_assignment.png",
    "vertical_management_performance": IMG_DIR / "11_vertical_management_performance.png",
    "vertical_management_new_customers": IMG_DIR / "12_vertical_management_new_customers.png",
    "vertical_management_potential": IMG_DIR / "13_vertical_management_potential.png",
    "vertical_management_drill_dialog": IMG_DIR / "14_vertical_management_drill_dialog.png",
    "alerts_top": IMG_DIR / "15_alerts_top.png",
    "alerts_scrolled": IMG_DIR / "16_alerts_scrolled.png",
    "alert_detail": IMG_DIR / "17_alert_detail.png",
    "alert_detail_script": IMG_DIR / "18_alert_detail_script.png",
    "analysis_risk": IMG_DIR / "19_analysis_risk.png",
    "analysis_cashflow": IMG_DIR / "20_analysis_cashflow.png",
    "analysis_recommend": IMG_DIR / "21_analysis_recommend.png",
    "analysis_script": IMG_DIR / "22_analysis_script.png",
    "skills_home": IMG_DIR / "23_skills_home.png",
    "skills_detail": IMG_DIR / "24_skills_detail.png",
    "skills_edit": IMG_DIR / "25_skills_edit.png",
    "skills_create": IMG_DIR / "26_skills_create.png",
    "qa_home": IMG_DIR / "27_qa_home.png",
    "qa_result": IMG_DIR / "28_qa_result.png",
    "qa_knowledge_base": IMG_DIR / "29_qa_knowledge_base.png",
    "workflow_list": IMG_DIR / "30_workflow_list.png",
    "workflow_editor": IMG_DIR / "31_workflow_editor.png",
    "workflow_run_panel": IMG_DIR / "32_workflow_run_panel.png",
    "tasks_home": IMG_DIR / "33_tasks_home.png",
    "tasks_create": IMG_DIR / "34_tasks_create.png",
}


def set_run_font(run, font_name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level=1):
    """Add a heading with Chinese font support."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, font_name="Microsoft YaHei", size=(18 if level == 1 else (14 if level == 2 else 12)), bold=True, color=(31, 78, 121) if level == 1 else None)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_paragraph(doc, text, bold=False, size=10.5, color=None, align=None, first_line_indent=0.0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=size, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    return p


def add_screenshot(doc, img_path, caption):
    """Insert a screenshot centered with caption below."""
    if not img_path.exists():
        add_paragraph(doc, f"[图片缺失: {img_path.name}]", color=(255, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    crun = cap.add_run(f"图：{caption}")
    set_run_font(crun, size=9, color=(89, 89, 89))


def add_toc(doc):
    """Insert a Word TOC field. User updates on first open."""
    add_heading_custom(doc, "目录", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)
    add_paragraph(doc, "（在 Word 中打开后，右键目录选择「更新域」即可刷新页码。）", size=9, color=(128, 128, 128))
    doc.add_page_break()


def add_page_number_footer(section):
    """Add a centered page number footer to the section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("第 ")
    set_run_font(run, size=9, color=(128, 128, 128))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = p.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_separate)
    r._r.append(fld_end)

    run2 = p.add_run(" 页 / 共 ")
    set_run_font(run2, size=9, color=(128, 128, 128))

    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = "NUMPAGES"
    fld_separate2 = OxmlElement("w:fldChar")
    fld_separate2.set(qn("w:fldCharType"), "separate")
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld_begin2)
    r2._r.append(instr2)
    r2._r.append(fld_separate2)
    r2._r.append(fld_end2)

    run3 = p.add_run(" 页")
    set_run_font(run3, size=9, color=(128, 128, 128))


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for level in range(1, 4):
        try:
            style = styles[f"Heading {level}"]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.bold = True
            style.font.color.rgb = RGBColor(31, 78, 121) if level == 1 else RGBColor(0, 0, 0)
        except Exception:
            pass


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 客户经营助手")
    set_run_font(run, font_name="Microsoft YaHei", size=32, bold=True, color=(31, 78, 121))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("操作手册")
    set_run_font(run, font_name="Microsoft YaHei", size=24, bold=True, color=(68, 114, 196))

    doc.add_paragraph()
    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org.add_run("龙湾农村商业银行 · 演示系统")
    set_run_font(run, font_name="Microsoft YaHei", size=14, color=(89, 89, 89))

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run(f"版本：v1.0    生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(run, font_name="Microsoft YaHei", size=11, color=(128, 128, 128))

    doc.add_page_break()


def add_overview(doc):
    add_heading_custom(doc, "第一章 系统概述", level=1)

    add_heading_custom(doc, "1.1 产品定位", level=2)
    add_paragraph(doc,
        "AI 客户经营助手是面向银行客户经理、支行管理人员打造的智能经营工具。 "
        "系统通过自然语言交互、预置业务模板和 Agent 工作流，将原本分散在多个业务系统中的客户筛选、 "
        "绩效跟踪、风险预警、客户画像等能力整合到一个统一界面中，帮助一线人员用更短的时间完成更精准的客户经营决策。")

    add_heading_custom(doc, "1.2 解决的核心痛点", level=2)
    add_bullet(doc, "客户数据分散：存款、贷款、征信、走访记录需要在多个系统间切换查询。")
    add_bullet(doc, "筛选效率低：手工拼 Excel 耗时长，条件稍复杂就容易出错。")
    add_bullet(doc, "预警滞后：存款/贷款到期、融资异动等机会与风险难以及时触达客户经理。")
    add_bullet(doc, "管理口径不统一：支行层面难以实时掌握每位客户经理的新增贡献与维护情况。")

    add_heading_custom(doc, "1.3 核心能力", level=2)
    add_bullet(doc, "自然语言任务：输入一句话即可生成客户清单、扫描预警、分析风险。")
    add_bullet(doc, "客群梳理：5 套预置模板 + 自定义条件，实时统计分布并导出 Excel/CSV。")
    add_bullet(doc, "垂直管理：客户经理绩效排名、支行名单导入、AI 自动分配。")
    add_bullet(doc, "业务预警：7 类预警主动推送，支持详情查看、话术生成与状态流转。")
    add_bullet(doc, "查询分析：输入客户姓名/编号/地址，一键生成 360° 画像与调查报告。")
    add_bullet(doc, "Skill Center：可加载/自定义 AI 行为规范的「技能片段」。")
    add_bullet(doc, "编排工作流：可视化节点组合多步 AI 分析流程。")

    add_heading_custom(doc, "1.4 适用角色与权限", level=2)
    add_bullet(doc, "客户经理（lixue）：重点关注客户清单、预警处理、客户画像。")
    add_bullet(doc, "支行负责人（zhoujianhua）：查看团队绩效、处理预警、审批分配。")
    add_bullet(doc, "分行管理员（admin）：全量菜单，包含渠道配置与数据源管理。")
    add_bullet(doc, "合规审计（compliance）：可访问审计日志，查看 AI 操作记录。")

    add_heading_custom(doc, "1.5 技术架构简介", level=2)
    add_paragraph(doc,
        "前端基于 Next.js 16 + React 19 + Tailwind CSS v4 构建，数据层使用 SQLite 本地存储， "
        "Agent 层支持真实 LLM（已对接 qwen-plus）与本地 Mock 双轨运行。 在演示环境或未配置 API 密钥时， "
        "系统会自动降级到本地 Mock，保证现场汇报稳定可控。")

    add_screenshot(doc, SCREENSHOTS["login"], "系统登录页与演示账号")
    add_screenshot(doc, SCREENSHOTS["dashboard_empty"], "登录后进入 AI 工作台首页")


def add_quickstart(doc):
    add_heading_custom(doc, "第二章 快速入门", level=1)

    add_heading_custom(doc, "2.1 访问与登录", level=2)
    add_numbered(doc, "在浏览器中打开系统地址（默认 http://localhost:3003）。")
    add_numbered(doc, "输入用户名与密码，演示环境默认账号为 admin / demo123。")
    add_numbered(doc, "点击「登录」后进入 AI 工作台。")

    add_heading_custom(doc, "2.2 界面布局说明", level=2)
    add_paragraph(doc, "系统采用经典的三栏式后台布局：")
    add_bullet(doc, "左侧 Sidebar：功能导航，包含 AI 工作台、客群梳理、垂直管理、业务预警、查询分析、问答助手、Skill Center、编排工作流、审计日志、渠道配置、数据源、定时任务。")
    add_bullet(doc, "顶部 Header：显示当前机构「龙湾农村商业银行」、演示环境标识、当前登录用户与退出按钮。")
    add_bullet(doc, "主内容区：根据所选菜单展示对应业务页面。")

    add_heading_custom(doc, "2.3 通用操作", level=2)
    add_bullet(doc, "快捷指令：AI 工作台底部提供 4 条常用 prompt，点击即可自动填入并执行。")
    add_bullet(doc, "筛选与导出：多数列表页支持关键词搜索、下拉筛选，并可通过右上角按钮导出 Excel 或 CSV。")
    add_bullet(doc, "卡片点击：业务预警、工作流等页面使用卡片式交互，点击后可在 Drawer/弹窗中查看详情。")


def add_workbench(doc):
    add_heading_custom(doc, "第三章 AI 工作台", level=1)

    add_heading_custom(doc, "3.1 功能价值", level=2)
    add_paragraph(doc,
        "AI 工作台是整个系统的入口。用户无需记忆复杂菜单，直接用自然语言描述业务需求， "
        "系统即可自动识别意图、调用工具、展示执行步骤并输出结构化结果。左侧会话栏支持新建会话， "
        "便于将不同业务问题分开展示。")

    add_heading_custom(doc, "3.2 快捷指令", level=2)
    add_paragraph(doc, "系统预置了 4 条高频快捷指令，点击即可一键执行：")
    add_bullet(doc, "梳理 高新区·锦园 中日均存款大于 10 万元的客户清单")
    add_bullet(doc, "扫描本月所有业务预警")
    add_bullet(doc, "分析 张明 的风险情况，并生成调查报告")
    add_bullet(doc, "统计各客户经理本月新增存款客户")

    add_heading_custom(doc, "3.3 典型对话场景", level=2)

    add_heading_custom(doc, "场景一：客户清单梳理", level=3)
    add_paragraph(doc, "点击快捷指令「梳理 高新区·锦园 中日均存款大于 10 万元的客户清单」，AI 会在对话中展示工具调用过程， "
        "并输出命中客户表格。表格右上角可直接导出 Excel 或 CSV。")
    add_screenshot(doc, SCREENSHOTS["dashboard_customer_segment"], "AI 工作台执行客户清单梳理")

    add_heading_custom(doc, "场景二：业务预警扫描", level=3)
    add_paragraph(doc, "点击快捷指令「扫描本月所有业务预警」，系统自动扫描并返回预警卡片列表， "
        "帮助客户经理快速掌握当前需要跟进的客户。")
    add_screenshot(doc, SCREENSHOTS["dashboard_business_alert"], "AI 工作台执行业务预警扫描")

    add_heading_custom(doc, "场景三：客户风险分析", level=3)
    add_paragraph(doc, "点击快捷指令「分析 张明 的风险情况，并生成调查报告」，AI 返回客户画像与调查报告预览， "
        "并支持跳转至「查询分析」页面查看完整画像。")
    add_screenshot(doc, SCREENSHOTS["dashboard_customer_analysis"], "AI 工作台执行客户风险分析")

    add_heading_custom(doc, "场景四：经理绩效统计", level=3)
    add_paragraph(doc, "点击快捷指令「统计各客户经理本月新增存款客户」，AI 返回经理绩效排名图表， "
        "便于支行负责人进行过程管理与绩效跟踪。")
    add_screenshot(doc, SCREENSHOTS["dashboard_vertical_management"], "AI 工作台执行经理绩效统计")


def add_customer_segments(doc):
    add_heading_custom(doc, "第四章 客群梳理", level=1)

    add_heading_custom(doc, "4.1 功能价值", level=2)
    add_paragraph(doc,
        "客群梳理帮助客户经理按存款、贷款、合同、征信等维度快速锁定目标客户。 "
        "系统提供 5 套预置模板，并允许在模板基础上继续微调小区/网格、金额阈值等条件。")

    add_heading_custom(doc, "4.2 预置模板", level=2)
    add_bullet(doc, "全量客户：不应用任何条件，作为筛选起点。")
    add_bullet(doc, "高日均存款客户：小区/网格 + 日均存款 ≥ 指定金额。")
    add_bullet(doc, "低贷高信客户：抵押贷款 > X 且 信用贷款 < Y。")
    add_bullet(doc, "有合同未用信：有效合同且当前用信为 0。")
    add_bullet(doc, "近期征信更新：近一年征信报告更新时间不为空。")
    add_bullet(doc, "无贷有贷（他行）：本行无贷，但他行有贷。")

    add_heading_custom(doc, "4.3 自定义筛选", level=2)
    add_paragraph(doc,
        "选择模板后，可在「自定义筛选器」中调整小区/网格、日均存款下限等条件，结果实时刷新。 "
        "页面中部提供风险等级、客群分类、日均存款区间的可视化分布。")
    add_screenshot(doc, SCREENSHOTS["customer_segments_top"], "客群梳理：预置模板、自定义筛选器与客群分布统计")

    add_heading_custom(doc, "4.4 客户清单与导出", level=2)
    add_paragraph(doc,
        "页面下方的客户清单支持关键词搜索、风险等级筛选、客户经理筛选、列排序、分页与多选。 "
        "右上角提供「导出全量 Excel」「CSV」按钮，便于下载后离线使用。")
    add_screenshot(doc, SCREENSHOTS["customer_segments_list"], "客群梳理：客户清单与导出功能")


def add_vertical_management(doc):
    add_heading_custom(doc, "第五章 客户经理垂直管理", level=1)

    add_heading_custom(doc, "5.1 功能价值", level=2)
    add_paragraph(doc,
        "垂直管理从支行视角出发，帮助管理者掌握每位客户经理的新增客户、新增存贷、维护得分等核心指标， "
        "并支持支行客户名单的导入与 AI 自动分配。")

    add_heading_custom(doc, "5.2 核心指标与贡献排名", level=2)
    add_paragraph(doc,
        "页面顶部展示本月新增客户合计、新增存款合计、新增贷款合计、维护得分均值四张核心指标卡。 "
        "中部图表展示本月客户经理贡献排名，可直观对比存款与贷款增量。")
    add_screenshot(doc, SCREENSHOTS["vertical_management_top"], "垂直管理：指标卡、贡献排名与名单导入")

    add_heading_custom(doc, "5.3 名单导入与 AI 自动分配", level=2)
    add_numbered(doc, "点击「导入支行客户清单」，系统读取演示文件「支行客户名单_2026_06.xlsx」。")
    add_numbered(doc, "查看总客户数、可分配客户数、AI 已匹配经理数、待人工确认数。")
    add_numbered(doc, "点击「AI 自动分配客户经理」，系统根据网格匹配、工作负荷、维护得分综合分配。")
    add_numbered(doc, "在「AI 自动分配结果」中查看分配详情，并可导出 Excel。")
    add_screenshot(doc, SCREENSHOTS["vertical_management_assignment"], "垂直管理：AI 自动分配结果详情")

    add_heading_custom(doc, "5.4 多维度绩效明细", level=2)
    add_paragraph(doc, "页面底部提供三个 Tab，满足不同管理视角：")
    add_bullet(doc, "经理绩效排名：按存款/贷款增量排序，前十名高亮显示。")
    add_bullet(doc, "本月新增存贷客户：按经理筛选的明细清单。")
    add_bullet(doc, "扩中客群贷款：各经理名下扩中客户新增贷款情况。")
    add_screenshot(doc, SCREENSHOTS["vertical_management_performance"], "垂直管理：经理绩效排名")
    add_screenshot(doc, SCREENSHOTS["vertical_management_new_customers"], "垂直管理：本月新增存贷客户")
    add_screenshot(doc, SCREENSHOTS["vertical_management_potential"], "垂直管理：扩中客群贷款统计")

    add_heading_custom(doc, "5.5 经理下钻 Dialog", level=2)
    add_paragraph(doc, "在经理绩效排名表格中点击「查看客户」，可弹出该经理今年新引入客户清单，支持导出 Excel。")
    add_screenshot(doc, SCREENSHOTS["vertical_management_drill_dialog"], "垂直管理：经理新引入客户下钻弹窗")


def add_alerts(doc):
    add_heading_custom(doc, "第六章 业务预警", level=1)

    add_heading_custom(doc, "6.1 功能价值", level=2)
    add_paragraph(doc,
        "业务预警主动扫描存款到期、贷款到期、融资增长、网格变动、支行异常等多源数据， "
        "将「被动等客户上门」转变为「主动找客户跟进」。")

    add_heading_custom(doc, "6.2 预警类型与筛选", level=2)
    add_paragraph(doc, "系统支持以下预警类型，可通过类型、严重度、状态、关键词进行筛选：")
    add_bullet(doc, "存款即将到期 / 贷款即将到期")
    add_bullet(doc, "融资额增长 / 融资额大幅上浮")
    add_bullet(doc, "新楼盘 / 网格变动 / 支行异常")
    add_screenshot(doc, SCREENSHOTS["alerts_top"], "业务预警：统计卡、筛选条件与预警卡片")
    add_screenshot(doc, SCREENSHOTS["alerts_scrolled"], "业务预警：向下滚动查看更多预警卡片")

    add_heading_custom(doc, "6.3 处理流程与话术", level=2)
    add_numbered(doc, "在预警列表中点击任意卡片，右侧弹出详情 Drawer。")
    add_numbered(doc, "查看触发原因、涉及客户/金额、AI 建议下一步、数据来源。")
    add_numbered(doc, "点击「生成联系话术」获取 AI 推荐沟通文本。")
    add_numbered(doc, "可选择「推送企业微信」「指派经理」「标记处理中」「标记已完成」等操作。")
    add_screenshot(doc, SCREENSHOTS["alert_detail"], "业务预警详情 Drawer：触发原因与 AI 建议")
    add_screenshot(doc, SCREENSHOTS["alert_detail_script"], "业务预警详情 Drawer：AI 沟通话术")


def add_analysis(doc):
    add_heading_custom(doc, "第七章 查询分析", level=1)

    add_heading_custom(doc, "7.1 功能价值", level=2)
    add_paragraph(doc,
        "查询分析实现「输入一个客户，输出一份完整调查报告」。 "
        "客户经理无需再分别登录核心系统、征信系统、反洗钱系统，系统会自动汇总并生成可复核的分析结论。")

    add_heading_custom(doc, "7.2 搜索方式", level=2)
    add_paragraph(doc, "支持输入客户姓名、客户编号、身份证号或地址关键词，例如「张明」「C001」「高新区·锦园」。")

    add_heading_custom(doc, "7.3 客户 360° 画像", level=2)
    add_paragraph(doc, "搜索结果以卡片形式展示客户基础信息、风险等级、所在网格、日均存款、授信余额、月均净流入等关键指标。 "
        "下方 Tab 可进一步查看：")
    add_bullet(doc, "风险与准入：风险信号、黑名单/限入名单/中介名单核查结果、AI 调查报告。")
    add_bullet(doc, "资金流分析：近 6 个月资金流入流出趋势、主要上下游关系。")
    add_bullet(doc, "产品推荐：存款/理财、贷款产品的匹配推荐。")
    add_bullet(doc, "走访记录：历史走访时间、渠道、经理、摘要。")
    add_bullet(doc, "营销话术：按营销、催收、续存、转介绍场景生成沟通文本。")
    add_screenshot(doc, SCREENSHOTS["analysis_risk"], "查询分析：风险信号与准入判断")
    add_screenshot(doc, SCREENSHOTS["analysis_cashflow"], "查询分析：资金流分析与上下游关系")
    add_screenshot(doc, SCREENSHOTS["analysis_recommend"], "查询分析：存款/理财与贷款产品推荐")
    add_screenshot(doc, SCREENSHOTS["analysis_script"], "查询分析：营销话术与沟通技巧")


def add_skills(doc):
    add_heading_custom(doc, "第八章 Skill Center", level=1)

    add_heading_custom(doc, "8.1 功能价值", level=2)
    add_paragraph(doc,
        "Skill 是一段注入 AI 上下文的行为规范。加载后，AI 的回答风格、关注重点、输出格式会真实改变。 "
        "例如加载「合规审查员」后，AI 每次输出结论都会自动附加合规提示。")

    add_heading_custom(doc, "8.2 内置 Skill 一览", level=2)
    add_bullet(doc, "保守风控视角：主动列出潜在风险点，不轻易给出「通过」结论。")
    add_bullet(doc, "合规审查员：自动附加合规提示与人工复核建议。")
    add_bullet(doc, "简报模式：将回答压缩到 3 句以内，适合移动端快速浏览。")
    add_bullet(doc, "营销话术优化：遵循 FABE 结构，语气亲切且有明确行动引导。")
    add_bullet(doc, "支行管理视角：关注团队绩效分布与资源分配。")
    add_bullet(doc, "催收专家：按五级分类给出差异化催收策略。")
    add_screenshot(doc, SCREENSHOTS["skills_home"], "Skill Center：内置与自定义 AI 行为规范卡片")

    add_heading_custom(doc, "8.3 Skill 详情与编辑", level=2)
    add_paragraph(doc, "点击卡片上的「查看」可打开详情 Dialog，查看该 Skill 注入的 Prompt 片段；点击「编辑」可在右侧 Sheet 中修改提示词。")
    add_screenshot(doc, SCREENSHOTS["skills_detail"], "Skill Center：Skill 详情 Dialog")
    add_screenshot(doc, SCREENSHOTS["skills_edit"], "Skill Center：编辑 Skill 提示词 Sheet")

    add_heading_custom(doc, "8.4 自定义 Skill", level=2)
    add_paragraph(doc, "点击右上角「新建 Skill」，填写名称、描述、分类和提示词内容，即可创建自定义行为规范。自定义内容保存在本地存储或后端数据库中。")
    add_screenshot(doc, SCREENSHOTS["skills_create"], "Skill Center：新建自定义 Skill Sheet")


def add_extensions(doc):
    add_heading_custom(doc, "第九章 扩展模块", level=1)

    add_heading_custom(doc, "9.1 问答助手", level=2)
    add_paragraph(doc,
        "问答助手用于查询贷款政策、利率、合规要求等内部知识。系统内置 15 条知识库， "
        "每条答案均标注文件来源，便于复核。")
    add_screenshot(doc, SCREENSHOTS["qa_home"], "问答助手：查询首页与示例问题")
    add_screenshot(doc, SCREENSHOTS["qa_result"], "问答助手：查询结果与文件来源")
    add_screenshot(doc, SCREENSHOTS["qa_knowledge_base"], "问答助手：内置知识库列表弹窗")

    add_heading_custom(doc, "9.2 编排工作流", level=2)
    add_paragraph(doc,
        "编排工作流提供可视化拖拽画布，可将「筛选客户 → 风险分析 → 生成话术」「预警扫描 → 条件分支 → 分级处理」等流程固化为可复用的 Agent 工作流。")
    add_screenshot(doc, SCREENSHOTS["workflow_list"], "编排工作流：预置模板与我的工作流列表")
    add_screenshot(doc, SCREENSHOTS["workflow_editor"], "编排工作流：可视化节点编辑器")
    add_screenshot(doc, SCREENSHOTS["workflow_run_panel"], "编排工作流：运行面板与日志")

    add_heading_custom(doc, "9.3 定时任务", level=2)
    add_paragraph(doc,
        "定时任务用于配置周期性 AI 任务，例如「每周一早 8 点扫描本周到期存款并推送给客户经理」。 "
        "支持按日/周/月配置执行计划，并查看历史执行记录。")
    add_screenshot(doc, SCREENSHOTS["tasks_home"], "定时任务：任务列表与空状态")
    add_screenshot(doc, SCREENSHOTS["tasks_create"], "定时任务：新建任务 Sheet")


def main():
    doc = Document()
    setup_styles(doc)

    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)
    add_page_number_footer(sections)

    add_cover(doc)
    add_toc(doc)
    add_overview(doc)
    add_quickstart(doc)
    add_workbench(doc)
    add_customer_segments(doc)
    add_vertical_management(doc)
    add_alerts(doc)
    add_analysis(doc)
    add_skills(doc)
    add_extensions(doc)

    doc.save(str(OUT_FILE))
    print(f"Manual generated: {OUT_FILE}")


if __name__ == "__main__":
    main()
