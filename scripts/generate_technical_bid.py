from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "generated/温州龙湾农商行AI智能体项目技术标书.docx"

BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
WHITE = RGBColor(255, 255, 255)


def set_font(run, name="微软雅黑", size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para(p, before=0, after=6, line=1.15, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5, color=DARK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.12, align=align)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.style.font.name = "微软雅黑"
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK)
    set_para(p, before={1: 16, 2: 10, 3: 6}.get(level, 6), after={1: 8, 2: 6, 3: 4}.get(level, 4), line=1.12)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_para(p, after=6, line=1.18)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, color=DARK)
    else:
        r = p.add_run(text)
        set_font(r, color=DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, after=4, line=1.15)
        r = p.add_run(item)
        set_font(r, color=DARK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para(p, after=4, line=1.15)
        r = p.add_run(item)
        set_font(r, color=DARK)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_cell_margins(table)
    for idx, h in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], h, bold=True, size=10.5, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            set_cell_text(cells[idx], str(val), size=10, color=DARK)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.35])
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    shade_cell(cell, "EEF6FC")
    cell.text = ""
    p = cell.paragraphs[0]
    set_para(p, after=4, line=1.15)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    set_para(p2, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color=DARK)
    doc.add_paragraph()


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("温州龙湾农商行 AI 智能体项目技术标书")
    set_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("技术响应文件")
    set_font(r, size=9, color=MUTED)


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("温州龙湾农商行")
    set_font(r, size=18, bold=True, color=BLUE)
    set_para(p, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 智能体项目")
    set_font(r, size=26, bold=True, color=BLUE)
    set_para(p, after=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("技术标书")
    set_font(r, size=24, bold=True, color=DARK)
    set_para(p, after=30)

    rows = [
        ("项目名称", "温州龙湾农商行 AI 智能体项目"),
        ("文件类型", "技术标书 / 技术响应文件"),
        ("适用场景", "客群梳理、垂直管理、业务预警、查询分析、知识问答"),
        ("编制依据", "项目申报书、演示 Demo、银行一线经营管理需求"),
        ("版本日期", "2026 年 5 月"),
    ]
    add_table(doc, ["项目", "内容"], rows, [1.6, 4.75], header_fill=LIGHT_BLUE)
    doc.add_page_break()


def toc_like(doc):
    add_heading(doc, "目录", 1)
    entries = [
        "一、项目理解与建设目标",
        "二、总体技术方案",
        "三、核心功能技术响应",
        "四、智能体能力设计",
        "五、数据架构与系统集成",
        "六、安全合规与运维保障",
        "七、实施计划、测试验收与交付物",
        "八、技术评分项建议",
    ]
    for e in entries:
        p = doc.add_paragraph()
        set_para(p, after=3, line=1.15)
        r = p.add_run(e)
        set_font(r, size=11.5, color=DARK)
    doc.add_page_break()


def build_content(doc):
    add_heading(doc, "一、项目理解与建设目标", 1)
    add_para(doc, "本项目面向温州龙湾农商行客户经理、支行管理人员及总行经营管理部门，建设以 AI 智能体为核心的数智化经营赋能平台。平台以自然语言交互为入口，以客户数据、业务规则、知识库和预警事件为基础，支持客户经营、网格管理、风险提示、绩效分析和政策问答等高频业务场景。")
    add_para(doc, "系统建设目标是将“客群梳理、垂直管理、业务预警、查询分析”四类需求转化为可操作、可追踪、可审计的数字化能力，使一线人员能够通过自然语言完成数据查询、名单生成、客户画像分析、风险预警处理和营销话术生成，满足银行精细化经营、穿透式管理和合规化运营要求。")
    add_callout(doc, "建设原则", "坚持业务可落地、数据可追溯、模型可管控、结果可解释、演示可验证、生产可扩展。现有 Demo 已验证自然语言指令、Agent 工具调用、结构化结果呈现、Excel/CSV 导出、预警处理和客户画像等关键路径，可作为后续生产化建设的原型基础。")
    add_heading(doc, "1.1 需求响应范围", 2)
    rows = [
        ("客群梳理", "支持按小区、网格、日均存款、抵押贷款、信用贷款、有效合同、未用信、他行有贷、征信更新等条件生成客户清单，并支持导出。"),
        ("垂直管理", "支持客户经理新引入客户、绩效维护人、本月新增存贷、标签客群贷款增长、排名与环比变化分析。"),
        ("业务预警", "支持存款到期、贷款到期、融资家数增长、融资金额上浮、新楼盘、网格变动、网点异常波动等预警识别与播报。"),
        ("查询分析", "支持客户身份检索、360 度画像、网格归属、走访记录、征信摘要、产品推荐、风险初判、政策问答和个性化话术生成。"),
    ]
    add_table(doc, ["业务方向", "技术响应说明"], rows, [1.25, 5.1])

    add_heading(doc, "二、总体技术方案", 1)
    add_para(doc, "平台采用“数据资源层—智能能力中台层—场景应用层”的分层架构。数据资源层负责整合核心系统、信贷系统、数据仓库、征信摘要、客户标签、产品政策、走访记录及预警事件；智能能力中台层封装语义解析、工具编排、规则引擎、RAG 知识问答、客户画像、报表导出等公共能力；场景应用层面向客户经理和管理人员提供统一工作台、管理驾驶舱、预警中心、查询分析和 Skill Center。")
    rows = [
        ("场景应用层", "客户经理智能工作台、客群梳理、垂直管理、业务预警、查询分析、合规问答、技能中心", "提供统一操作入口与可视化结果。"),
        ("智能能力中台层", "Agent 编排、语义解析、规则引擎、RAG 问答、客户画像、产品推荐、话术生成、导出服务", "将用户意图转化为可执行工具调用。"),
        ("数据资源层", "客户主数据、账户与存贷数据、征信摘要、走访记录、预警事件、产品政策、知识文档", "形成客户中心化数据模型和标签体系。"),
        ("基础设施层", "微服务、API 网关、PostgreSQL、MPP 分析库、向量数据库、对象存储、日志审计", "支撑生产部署、扩展和运维。"),
    ]
    add_table(doc, ["架构层级", "组成内容", "作用"], rows, [1.15, 3.15, 2.05])
    add_heading(doc, "2.1 Demo 技术基线与生产化演进路线", 2)
    add_bullets(doc, [
        "现有演示 Demo 采用 Next.js App Router + Node.js API Routes 作为前后端一体化原型架构，后端能力包括 SSE 流式接口、Agent 工具编排、本地 SQLite 数据服务、Excel/CSV 导出服务和企业微信推送模拟接口。",
        "Demo 中 AI Agent 采用“自然语言理解—工具调用—结构化结果返回”的方式运行，已验证客户筛选、预警扫描、客户画像、报告生成、数据导出、话术生成和知识问答等关键链路。",
        "正式生产建设阶段建议在保留 Demo 业务交互和 Agent 编排模式的基础上，结合银行现有技术规范演进为可独立部署的服务化架构。后端可采用 Spring Cloud、Java 微服务或行内已有微服务框架，将客户服务、规则服务、预警服务、知识问答服务、报表导出服务、权限审计服务等模块拆分。",
        "AI 能力采用国产大模型或 OpenAI-Compatible 模型接口，支持 DeepSeek、通义千问、私有化模型网关等接入方式，避免单一模型绑定。",
        "知识问答采用 RAG 架构，使用向量数据库存储政策制度、产品手册、合规要求和操作规范，答案必须关联来源文件。",
        "分析查询场景可引入 MPP 分析数据库或数据仓库加速层，对常用指标进行预计算，满足复杂条件组合下的查询性能要求。",
        "前端采用现代化 Web 技术构建响应式工作台，现有 Demo 已基于 Next.js、React、流式交互和结构化组件验证用户体验；生产阶段可继续沿用该交互形态，也可按行内前端技术规范迁移实现。",
    ])

    add_heading(doc, "三、核心功能技术响应", 1)
    add_heading(doc, "3.1 客群智能梳理", 2)
    add_para(doc, "系统支持客户经理使用业务口语直接提出客群筛选需求，例如“梳理某小区中日均存款大于 1 万元的客户清单”“梳理我行无贷、他行有贷客户清单”。智能体解析用户意图后，将自然语言映射为标准化查询条件，并调用客户查询工具获取结果。")
    add_bullets(doc, [
        "支持多条件组合筛选：小区、网格、客户标签、存款余额、贷款余额、合同状态、用信状态、征信更新时间等。",
        "支持结果字段标准化展示：客户名称、证件脱敏号、联系方式脱敏号、联系地址、所属网格、客户经理、存贷信息等。",
        "支持一键导出 Excel 或 CSV，满足名单下发、客户跟进和经营分析需要。",
        "支持预置模板与自定义查询并存，兼顾固定经营动作和临时分析需求。",
    ])
    add_heading(doc, "3.2 垂直管理", 2)
    add_para(doc, "系统支持支行和总行对客户经理经营过程进行穿透式管理。通过导入客户清单、维护网格规则、识别客户归属和绩效维护人，系统可自动生成客户经理维度的新增客户、新增存贷、标签客群拓展、排名及环比变化。")
    add_bullets(doc, [
        "支持依据“机构—网格—客户经理”规则自动分配客户，减少人工拆分和漏分。",
        "支持客户经理月度绩效看板，展示新增客户数、新增存款、新增贷款、维护得分和环比变化。",
        "支持标注贷款增量前十、存款增长排名、重点标签客户拓展情况，满足管理驾驶舱需求。",
        "支持从全行、支行、网格、客户经理多层级钻取分析。",
    ])
    add_heading(doc, "3.3 业务预警", 2)
    add_para(doc, "系统支持围绕客户生命周期和经营管理指标建立主动预警机制。对于存款到期、贷款到期、融资家数增长、融资金额异常上浮、新楼盘、网格变动、网点存贷款异常波动等事件，系统可根据规则自动识别并生成处理建议。")
    add_bullets(doc, [
        "支持固定周期预警，如存款到期、贷款到期提前半个月提醒客户经理对接。",
        "支持规则类异常预警，如融资金额增加 30 万元以上、融资家数增长、网点指标非正常波动等。",
        "支持预警等级、状态流转、详情查看、处理建议和话术生成。",
        "支持与企业微信等协同平台对接，将风险或机会线索推送至对应客户经理。",
    ])
    add_heading(doc, "3.4 查询分析", 2)
    add_para(doc, "系统支持通过客户姓名、身份证号片段、手机号片段、客户编号等关键信息查询客户 360 度画像，集中展示客户基础信息、资产负债、网格归属、走访记录、征信摘要、风险信号、准入结论和产品推荐。")
    add_bullets(doc, [
        "支持客户全景视图，帮助客户经理快速掌握客户基本情况和经营价值。",
        "支持结合风险偏好、资金流动性、存贷款记录和征信摘要生成存款或贷款产品推荐。",
        "支持根据历史数据和五级分类生成个性化营销、催收、续存或转介绍话术。",
        "支持 AI 实时问答助手，回答贷款政策、利率、准入、合规要求等问题，并返回来源文件。",
    ])

    add_heading(doc, "四、智能体能力设计", 1)
    add_para(doc, "智能体采用“意图识别—工具选择—参数抽取—工具执行—结果总结—结构化展示”的运行机制。模型不直接编造业务结果，而是根据用户指令调用受控工具，由工具从数据库、规则引擎或知识库中获取确定性结果，再由模型生成简洁业务结论。")
    rows = [
        ("意图识别", "识别用户属于客群梳理、垂直管理、预警扫描、客户分析、报告生成、数据导出或知识问答。"),
        ("工具调用", "调用 filterCustomers、scanAlerts、analyzeCustomer、generateReport、exportData、queryDatabase 等受控工具。"),
        ("结构化结果", "根据工具返回类型展示表格、预警卡片、客户画像、报告文本、文件下载链接等。"),
        ("技能加载", "支持风控视角、合规审查、简报模式、营销话术、支行管理视角等可插拔技能提示。"),
        ("流式交互", "支持执行步骤和回答内容流式展示，提高一线人员等待体验和任务透明度。"),
    ]
    add_table(doc, ["能力环节", "技术说明"], rows, [1.25, 5.1])
    add_heading(doc, "4.1 RAG 知识问答", 2)
    add_para(doc, "知识问答模块采用检索增强生成技术。系统先在行内知识库中检索相关制度、产品说明、利率政策、合规要求和操作手册，再将检索结果提供给大模型生成回答。回答内容应附带来源文件、条款或文档标题，满足可追溯和可复核要求。")
    add_heading(doc, "4.2 模型安全边界", 2)
    add_bullets(doc, [
        "模型仅负责意图理解、参数抽取和自然语言表达，关键业务数据由系统工具和数据库返回。",
        "涉及授信、风险评级、催收和合规判断时，系统提示必须标明 AI 辅助属性，并要求客户经理人工复核。",
        "敏感字段默认脱敏展示，严禁输出完整身份证号、手机号、银行卡号等敏感个人信息。",
        "所有工具调用、用户提问、模型回答、导出动作和预警处理均应记录审计日志。",
    ])

    add_heading(doc, "五、数据架构与系统集成", 1)
    add_para(doc, "系统建设应围绕客户中心化数据模型展开，统一整合客户基础信息、账户资产、贷款余额、合同信息、用信情况、征信摘要、走访记录、网格归属、客户经理维护关系、业务预警和产品知识。")
    rows = [
        ("关系型数据库", "PostgreSQL 或同等级数据库", "存储系统配置、用户权限、客户关系、规则配置、预警事件、操作日志。"),
        ("分析型数据库", "MPP / 数据仓库 / 数据集市", "支撑多表关联、复杂筛选、指标聚合和经营看板。"),
        ("向量数据库", "Milvus、pgvector 或同等级组件", "存储制度文档、产品手册、合规知识的向量索引。"),
        ("对象存储", "行内文件服务或对象存储", "存储导出文件、知识文档、测试报告和模型评估材料。"),
        ("消息与集成", "API 网关、消息队列、企业微信接口", "支持预警播报、系统集成和异步任务处理。"),
    ]
    add_table(doc, ["数据/集成组件", "建议选型", "用途"], rows, [1.3, 2.0, 3.05])
    add_heading(doc, "5.1 数据质量与权限控制", 2)
    add_bullets(doc, [
        "建立客户主数据唯一标识，避免同名客户、跨网格客户和历史归属变更导致结果不一致。",
        "按总行、支行、网格、客户经理角色控制数据可见范围，确保客户经理仅查看授权客户和授权字段。",
        "对导出操作设置权限、审批或水印策略，避免客户清单被非授权下载。",
        "建立数据口径管理机制，对日均存款、新增贷款、有效客户、扩中客群等指标统一定义。",
    ])

    add_heading(doc, "六、安全合规与运维保障", 1)
    add_para(doc, "平台建设应满足银行信息安全、数据安全和合规审计要求。系统需支持私有化部署或专有网络部署，大模型访问需通过模型网关或行内授权通道进行统一管理。")
    add_bullets(doc, [
        "身份认证：支持统一身份认证、角色权限、菜单权限、数据权限和接口权限控制。",
        "数据安全：敏感字段脱敏展示，导出文件可记录操作人、时间、查询条件和下载来源。",
        "模型安全：支持提示词模板管理、敏感词过滤、答案来源约束、人工复核提示和模型调用审计。",
        "运行监控：支持接口耗时、模型调用成功率、工具调用失败率、预警推送状态、导出任务状态监控。",
        "高可用设计：关键服务支持水平扩展，异步任务和预警推送支持失败重试。",
    ])
    add_heading(doc, "6.1 性能指标响应", 2)
    rows = [
        ("常规查询响应", "单表或简单聚合查询平均响应时间不高于 3 秒。"),
        ("复杂查询响应", "多表关联、复杂条件组合查询平均响应时间不高于 10 秒。"),
        ("任务成功率", "有效用户提问的智能体任务成功解决率不低于 85%。"),
        ("语义解析准确率", "封闭测试集自然语言到查询条件转换准确率不低于 85%。"),
        ("导出能力", "支持客户、经理、预警等清单导出为 Excel 或 CSV，导出字段可配置。"),
    ]
    add_table(doc, ["指标项", "响应内容"], rows, [1.6, 4.75])

    add_heading(doc, "七、实施计划、测试验收与交付物", 1)
    add_para(doc, "项目实施建议采用“原型验证—接口联调—试点运行—优化推广”的阶段化方式推进。现有 Demo 可作为需求澄清和用户体验验证基础，后续围绕真实数据接入、权限控制、知识库建设和生产化部署逐步完善。")
    rows = [
        ("第一阶段：需求确认与方案细化", "梳理数据源、指标口径、权限范围、预警规则、知识文档目录和验收样例。", "需求规格说明、详细设计方案。"),
        ("第二阶段：核心能力开发", "完成客群梳理、垂直管理、预警中心、查询分析、知识问答、导出服务等核心模块。", "可联调系统、接口文档、测试用例。"),
        ("第三阶段：数据接入与试点运行", "接入试点数据，完成权限配置、规则配置、知识库导入和用户试用反馈。", "试点报告、问题清单、优化版本。"),
        ("第四阶段：验收交付与培训", "完成性能测试、功能验收、安全检查、用户培训和运维交接。", "验收报告、部署说明、用户手册、培训材料。"),
    ]
    add_table(doc, ["阶段", "主要工作", "成果物"], rows, [1.7, 3.0, 1.65])
    add_heading(doc, "7.1 测试验收建议", 2)
    add_numbered(doc, [
        "功能验收：按功能需求清单逐项验证自然语言查询、结构化展示、清单导出、预警处理、客户画像、知识问答和话术生成。",
        "准确性验收：建立封闭测试集，对语义解析、查询结果、知识问答来源、产品推荐逻辑进行抽样核验。",
        "性能验收：对常规查询、复杂查询、并发访问、导出任务和预警推送进行压力测试。",
        "安全验收：检查权限隔离、敏感字段脱敏、审计日志、导出控制和模型输出合规提示。",
    ])

    add_heading(doc, "八、技术评分项建议", 1)
    add_para(doc, "以下评分项建议用于技术标评审，满分 100 分。评分设置以项目实际落地能力为导向，避免设置过高或排他性门槛，同时兼顾智能体效果、银行业务适配、数据安全和交付保障。")
    rows = [
        (
            "需求理解与业务适配",
            "10",
            "评分说明：理解客群梳理、垂直管理、业务预警、查询分析四类需求；能结合银行客户经理和管理人员工作流程说明方案。证明方式：提供需求响应矩阵。扣分规则：缺少一类核心需求响应的，每项扣 1-2 分；无法对应银行实际岗位流程的，酌情扣 1-3 分。",
        ),
        (
            "总体技术架构",
            "15",
            "评分说明：架构分层清晰，支持 Demo 原型架构与生产化服务架构衔接，覆盖数据资源、AI 能力中台、场景应用和安全运维。证明方式：提供总体架构设计图。扣分规则：架构层级不完整的，扣 2-5 分；未说明 Demo 与生产化演进关系的，扣 1-3 分；关键组件职责不清的，扣 1-3 分。",
        ),
        (
            "核心功能覆盖度",
            "20",
            "评分说明：对功能需求清单中主要场景有明确响应，包括客户筛选、名单导出、绩效分析、预警播报、客户画像、知识问答和话术生成。证明方式：提供功能页面截图。扣分规则：每缺少一项主要业务闭环扣 2-3 分；仅文字描述、无可验证页面或原型的，酌情扣 2-5 分。",
        ),
        (
            "智能体与 RAG 能力",
            "15",
            "评分说明：具备自然语言意图识别、工具调用、结构化输出、知识检索增强、答案来源追溯和人工复核提示能力。证明方式：提供关键代码片段或接口调用示例。扣分规则：未体现工具调用链路的，扣 3-5 分；未体现知识来源追溯的，扣 2-4 分；缺少人工复核或合规提示机制的，扣 1-3 分。",
        ),
        (
            "数据治理与系统集成",
            "10",
            "评分说明：说明客户中心化数据模型、指标口径、数据质量、数据库选型、企业微信或行内系统集成方式。证明方式：提供数据模型或系统集成关系图。扣分规则：未说明核心数据实体和指标口径的，扣 2-4 分；未说明外部系统集成方式的，扣 1-3 分；数据质量控制描述不足的，扣 1-2 分。",
        ),
        (
            "安全合规与审计",
            "10",
            "评分说明：具备权限控制、敏感信息脱敏、导出审计、模型调用审计、合规提示和日志留存方案。证明方式：提供安全控制清单。扣分规则：缺少权限或脱敏设计的，扣 2-4 分；缺少导出和模型调用审计的，扣 2-3 分；合规提示和日志留存说明不足的，扣 1-3 分。",
        ),
        (
            "性能与稳定性",
            "8",
            "评分说明：响应指标合理，支持常规查询 ≤3 秒、复杂查询 ≤10 秒、失败重试、监控告警和高可用部署。证明方式：提供测试报告或压测结果摘要。扣分规则：未提供性能指标说明的，扣 2-3 分；未说明失败重试和监控告警的，扣 1-3 分；高可用部署描述不足的，扣 1-2 分。",
        ),
        (
            "实施交付与培训",
            "7",
            "评分说明：实施计划清晰，交付物完整，包含测试、验收、培训、运维交接和试点推广安排。证明方式：提供实施计划甘特图或阶段计划表。扣分规则：阶段计划不清晰的，扣 1-2 分；交付物不完整的，扣 1-3 分；缺少培训或运维交接安排的，扣 1-2 分。",
        ),
        (
            "演示与原型验证",
            "5",
            "评分说明：能够提供可运行 Demo 或原型，演示主要业务闭环，便于银行用户提前验证操作体验。证明方式：提供可访问 Demo 地址或现场演示截图。扣分规则：无法提供可运行原型的，扣 3-5 分；仅能演示单点功能、无法形成业务闭环的，扣 1-3 分。",
        ),
    ]
    add_table(doc, ["评分项", "分值", "评分说明"], rows, [1.55, 0.6, 4.2])
    add_callout(doc, "评分项设置说明", "上述评分项强调可落地能力和银行业务匹配度，不以特定厂商、特定模型或特定数据库作为排他条件。投标人可采用同等级、安全可控、满足性能与合规要求的技术组件进行响应。")

    add_heading(doc, "附：Demo 已验证能力对应关系", 1)
    rows = [
        ("AI 工作台", "自然语言输入、执行步骤 Timeline、流式回答、多会话本地保存、技能加载。"),
        ("Agent 工具链", "客户筛选、预警扫描、客户画像、调查报告、数据导出、话术生成、数据库查询。"),
        ("客群梳理", "预置模板、自定义条件、表格搜索排序分页、Excel/CSV 导出。"),
        ("垂直管理", "客户经理排名、月度绩效、导入清单、AI 自动分配演示。"),
        ("业务预警", "预警卡片、详情抽屉、状态流转、处理建议、企业微信推送模拟。"),
        ("查询分析", "客户 360 度画像、风险信号、准入结果、现金流分析、产品推荐、报告和话术。"),
        ("知识问答", "贷款政策、合规要求、利率调整等问答，并返回来源线索。"),
    ]
    add_table(doc, ["Demo 模块", "已验证内容"], rows, [1.35, 5.0])


def main():
    doc = Document()
    setup_document(doc)
    cover(doc)
    toc_like(doc)
    build_content(doc)
    doc.core_properties.title = "温州龙湾农商行AI智能体项目技术标书"
    doc.core_properties.subject = "AI智能体项目技术响应文件"
    doc.core_properties.author = "中科视语（北京）科技有限公司"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
